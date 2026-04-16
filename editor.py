"""
śikṣāmitra - Sanskrit Editor
PyQt6 + Flask hybrid application

This module provides the desktop application shell using PyQt6's WebEngineView
for rendering the web UI, with Flask serving the backend API.
"""

import os
import sys

# Suppress Qt DirectWrite font warnings (8514oem and other legacy fonts)
os.environ["QT_LOGGING_RULES"] = "qt.qpa.fonts.warning=false"

import json
import threading
import base64
import mimetypes
import shutil
import re
import socket
import time
import urllib.parse
import urllib.request
import zlib
import lzma
from datetime import datetime
import logging
try:
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from flask import Flask, request, jsonify, send_from_directory, Response
from bs4 import BeautifulSoup
from werkzeug.serving import make_server

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QFileDialog,
    QMessageBox, QSplashScreen, QLabel
)
from PyQt6.QtWebEngineWidgets import QWebEngineView
from PyQt6.QtWebChannel import QWebChannel
from PyQt6.QtWebEngineCore import QWebEngineSettings
from PyQt6.QtCore import (
    QUrl, Qt, QObject, pyqtSlot, pyqtSignal, QTimer, QSize
)
from PyQt6.QtGui import QIcon, QPixmap, QFont, QColor

# --- Configuration ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DOCS_DIR = os.path.join(BASE_DIR, 'documents')
MEDIA_DIR = os.path.join(BASE_DIR, 'media')
CACHE_DIR = os.path.join(BASE_DIR, 'cache')
LIBRARY_DIR = os.path.join(BASE_DIR, 'Library')
PREFS_FILE = os.path.join(BASE_DIR, 'preferences.json')
UNTITLED_PATH = os.path.join(CACHE_DIR, '_untitled.html')
SESSION_STATE_PATH = os.path.join(CACHE_DIR, '_last_session.json')
ABOUT_FILE = os.path.join(BASE_DIR, 'about.html')
LOGO_PATH = os.path.join(BASE_DIR, 'logo.png')
ICON_PATH = os.path.join(BASE_DIR, 'icon.ico')

PREFS_LOCK = threading.Lock()
SESSION_LOCK = threading.Lock()

# Ensure directories exist
for d in [DOCS_DIR, MEDIA_DIR, CACHE_DIR, LIBRARY_DIR]:
    os.makedirs(d, exist_ok=True)

# Suppress Flask/Werkzeug logging
logging.getLogger('werkzeug').setLevel(logging.ERROR)
logging.getLogger('werkzeug').disabled = True


# --- .smdoc Compression Helpers ---
# SMDI = lzma (current, best compression)
# SMDC = zlib level-9 (legacy, read-only)
SMDOC_MAGIC_LZMA = b'SMDI'   # lzma — current format
SMDOC_MAGIC_ZLIB = b'SMDC'   # zlib — legacy format (read support only)

def compress_smdoc(content: str) -> bytes:
    """Compress .smdoc content using lzma (best lossless compression)."""
    content_bytes = content.encode('utf-8')
    compressed = lzma.compress(content_bytes, preset=9 | lzma.PRESET_EXTREME)
    return SMDOC_MAGIC_LZMA + compressed

def decompress_smdoc(data: bytes) -> str:
    """Decompress .smdoc content. Supports lzma (SMDI), zlib (SMDC), and plain text."""
    if data.startswith(SMDOC_MAGIC_LZMA):
        return lzma.decompress(data[len(SMDOC_MAGIC_LZMA):]).decode('utf-8')
    elif data.startswith(SMDOC_MAGIC_ZLIB):
        return zlib.decompress(data[len(SMDOC_MAGIC_ZLIB):]).decode('utf-8')
    else:
        # Legacy plain-text .smdoc (no compression)
        return data.decode('utf-8')

def is_smdoc_file(filepath: str) -> bool:
    """Check if a file path is an .smdoc file."""
    return filepath.lower().endswith('.smdoc')


# --- .docx Import ---

# Map Word paragraph styles → editor paragraph CSS classes.
# Keys are matched exactly first, then by startswith prefix.
# Empty string '' = body text (plain <p>, no class).
_DOCX_PARA_STYLE_MAP = {
    'Title':         'ql-doc-title',
    'Subtitle':      'ql-doc-subtitle',
    'Heading 1':     'ql-doc-section',    # chapter/section header — large bold
    'Heading 2':     'ql-doc-subsection', # sub-section header
    'Heading 3':     'ql-doc-subsection',
    'Heading 4':     'ql-doc-subsection',
    'Prijevod':      'ql-doc-translation',
    'Translation':   'ql-doc-translation',
    # Primary mantra/text styles — plain body text (script detected separately)
    'Normal':        '',
    'Translit':      '',   # the "Translit" style used in IAST docs
    'Devanagari':    '',   # Devanagari text paragraphs
    'Telugu':        '',
    'Tamil':         '',
    'Kannada':       '',
    'Malayalam':     '',
    'Default':       '',
    'Body Text':     '',
}

# Map Word character styles → editor inline formatting
# Values: (css_class, wrapper_type)
#   wrapper_type: 'span' = <span class="...">
_DOCX_CHAR_STYLE_MAP = {
    # Holdings (from w:bdr borders or legacy named styles in IAST docs)
    'Holding':          ('ql-holding-short', 'span'),
    '2Holding':         ('ql-holding-long',  'span'),
    # Vedic accent marks
    'Svara':            ('ql-svara-char',    'span'),
    'Virama':           ('ql-svara-char',    'span'),
    # Superscript insertions (g in jñ, ś from ḥ, etc.) in IAST docs
    'Anusvara':         ('ql-change-style',  'span'),
    # Commentary runs
    'Comment':          ('ql-comment-style', 'span'),
    # Pause marks used in Devanagari/Indic sahasranāma docs:
    #   Long   = single vertical bar = short pause (|)
    #   Longer = double vertical bar = long pause  (||)
    'Long':             ('ql-short-pause',   'span'),
    'Longer':           ('ql-long-pause',    'span'),
}

# Maps IAST annotation characters → Devanagari equivalents for Insert/change-style
# runs inside Devanagari paragraphs.  The Veda Union DOCX format uses Latin/IAST
# chars as annotations even in Indic-script documents; we convert them on import
# so the editor shows native-script annotations.
# Consonants that cluster with the following char use virama (U+094D = ्).
# Vowels use the independent form (e.g. उ, not the mātrā ु) for clarity.
# Unknown chars pass through unchanged (identity) — so re-imported Devanagari
# annotations survive a second import without corruption.
_IAST_TO_DEV_ANNOTATION = {
    # Special Vedic insertions
    'g':  'ग',    # g in jñ → jgñ
    'u':  'उ',    # u in sv → suv / vy → vuy
    'i':  'इ',    # i insertion
    # Visarga transformation results (inline, half-form to cluster with next consonant)
    'ś':  'श\u094d',   # ḥ → ś before c/ch
    'ṣ':  'ष\u094d',   # ḥ → ṣ before ṭ/ṭh
    's':  'स\u094d',   # ḥ → s before t/th
    'r':  'र',          # ḥ → r before voiced (no virama — pre-vocalic)
    # Anusvara transformation results (half-form consonants)
    'ṅ':  'ङ\u094d',   # ṁ → ṅ before ka-varga
    'ñ':  'ञ\u094d',   # ṁ → ñ before ca-varga
    'ṇ':  'ण\u094d',   # ṁ → ṇ before ṭa-varga
    'n':  'न\u094d',   # ṁ → n before ta-varga
    'm':  'म\u094d',   # ṁ → m before pa-varga
}


# Pre-compiled regexes for _detect_script (C-speed instead of Python char loop)
_RE_DEVANAGARI = re.compile(r'[\u0900-\u097F\uA8E0-\uA8FF]')
_RE_TELUGU     = re.compile(r'[\u0C00-\u0C7F]')
_RE_TAMIL      = re.compile(r'[\u0B80-\u0BFF]')
_RE_KANNADA    = re.compile(r'[\u0C80-\u0CFF]')
_RE_MALAYALAM  = re.compile(r'[\u0D00-\u0D7F]')
_RE_LATIN      = re.compile(r'[\u0041-\u007A\u00C0-\u024F\u0300-\u036F]')


def _detect_script(text):
    """Detect the dominant Unicode script of text.

    Returns one of: 'devanagari', 'telugu', 'tamil', 'kannada', 'malayalam',
    'latin' (covers IAST), or None (empty / no dominant script).

    Uses compiled regexes (C-speed) rather than a Python character loop.
    """
    if not text:
        return None
    counts = {
        'devanagari': len(_RE_DEVANAGARI.findall(text)),
        'telugu':     len(_RE_TELUGU.findall(text)),
        'tamil':      len(_RE_TAMIL.findall(text)),
        'kannada':    len(_RE_KANNADA.findall(text)),
        'malayalam':  len(_RE_MALAYALAM.findall(text)),
        'latin':      len(_RE_LATIN.findall(text)),
    }
    total = sum(counts.values())
    if total == 0:
        return None
    dominant, count = max(counts.items(), key=lambda kv: kv[1])
    if count / total >= 0.25:
        return dominant
    return None


def _para_style_class(style_name):
    """Map a Word paragraph style name to an editor CSS class.

    Tries exact match first, then prefix match (e.g. 'Heading 1 Char' → 'Heading 1').
    Returns empty string for body-text styles.
    """
    if style_name in _DOCX_PARA_STYLE_MAP:
        return _DOCX_PARA_STYLE_MAP[style_name]
    # Prefix match
    for key, val in _DOCX_PARA_STYLE_MAP.items():
        if style_name.startswith(key):
            return val
    return None   # None = unknown style, treat as body text


def _html_escape(text):
    """Escape HTML special characters and preserve non-breaking spaces."""
    return (text
        .replace('&', '&amp;')
        .replace('<', '&lt;')
        .replace('>', '&gt;')
        .replace('\xa0', '\xa0')  # preserve non-breaking spaces as-is (Unicode)
    )


def _run_has_border(run):
    """Check if a run has a character border (holding)."""
    rPr = run._element.find(qn('w:rPr'))
    if rPr is not None:
        bdr = rPr.find(qn('w:bdr'))
        if bdr is not None:
            return True
    return False


def _run_border_size(run):
    """Get the border size of a run (to distinguish short vs long holding)."""
    rPr = run._element.find(qn('w:rPr'))
    if rPr is not None:
        bdr = rPr.find(qn('w:bdr'))
        if bdr is not None:
            sz = bdr.get(qn('w:sz'))
            if sz:
                return int(sz)
    return 0


def _detect_title_paragraph(para, idx):
    """Detect if a Normal-styled paragraph is actually a title based on formatting."""
    if idx > 2:
        return False
    # Check if any run has large font size (>= 24pt)
    for run in para.runs:
        if run.font.size and run.font.size >= 304800:  # 24pt in EMU
            return True
    # Check center alignment
    if para.alignment and para.alignment == 1:  # WD_ALIGN_PARAGRAPH.CENTER = 1
        return True
    return False


def _para_has_hanging_indent(para):
    """Check if paragraph (or its style) has a hanging indent."""
    pf = para.paragraph_format
    left = pf.left_indent
    first = pf.first_line_indent

    if left is None and para.style:
        left = para.style.paragraph_format.left_indent
    if first is None and para.style:
        first = para.style.paragraph_format.first_line_indent

    return bool(left and first and left > 0 and first < 0)


def convert_docx_to_html(filepath):
    """Convert a .docx file to Quill-compatible HTML for the editor.

    Maps Word styles to śikṣāmitra editor CSS classes:
    - Paragraph styles → ql-doc-title, ql-doc-section, etc.
    - Script detection → ql-script-devanagari / ql-script-telugu / ql-script-tamil
    - Character styles → ql-holding-short, ql-holding-long, ql-svara-char, ql-change-style
    - Run formatting → bold, italic, underline, superscript, font-size for 'Insert' runs
    """
    if not HAS_DOCX:
        raise ImportError("python-docx is required for Word import")

    doc = DocxDocument(filepath)
    html_parts = []

    for idx, para in enumerate(doc.paragraphs):
        # Skip completely empty paragraphs (spacers between sections)
        if not para.text.strip() and not para.runs:
            continue

        # Determine paragraph class from style name
        style_name = para.style.name if para.style else 'Normal'
        para_class = _para_style_class(style_name)

        # Unknown style (None) → treat as body text
        if para_class is None:
            para_class = ''

        # For very first Normal-looking paragraphs that look like a title, promote
        if not para_class and _detect_title_paragraph(para, idx):
            para_class = 'ql-doc-title'

        # Detect Unicode script of this paragraph's text
        script = _detect_script(para.text)
        script_class = f'ql-script-{script}' if script and script != 'latin' else ''

        # Build the inner HTML from runs
        hanging = _para_has_hanging_indent(para)
        inner_html = _convert_runs_to_html(para.runs, hanging_indent=hanging, para_script=script)

        # If paragraph is empty after conversion, add <br>
        if not inner_html.strip():
            inner_html = '<br>'

        # Build the <p> tag — combine paragraph class + script class
        classes = ' '.join(c for c in [para_class, script_class] if c)
        if classes:
            html_parts.append(f'<p class="{classes}">{inner_html}</p>')
        else:
            html_parts.append(f'<p>{inner_html}</p>')

    return '\n'.join(html_parts)


def _format_run_text(escaped_text, run, para_script=None):
    """Apply inline formatting to an already-escaped text fragment based on its run.

    para_script — the dominant script of the paragraph ('devanagari', 'telugu', …,
    'latin', or None).  Used to convert IAST annotation characters to their native-
    script equivalents when the paragraph is written in an Indic script.
    """
    char_style = None
    if run.style and run.style.name:
        char_style = run.style.name

    has_border = _run_has_border(run)
    border_sz = _run_border_size(run) if has_border else 0

    result = escaped_text

    # ── Detect Insert / annotation runs ──────────────────────────────────────
    # These are Vedic insertions (g in jñ, u in sv, ś from ḥ visarga, ṅ from anusvara …).
    # Detection priority:
    #   1. Named 'Insert'/'insert' character style
    #   2. Explicit small font (≤ 9pt) with no recognised char style
    #   3. Explicit superscript flag with no recognised char style
    cs_lower = char_style.lower() if char_style else ''
    is_insert = cs_lower.startswith('insert')
    if not is_insert and not (char_style and char_style in _DOCX_CHAR_STYLE_MAP):
        try:
            if run.font.size is not None and run.font.size.pt is not None and run.font.size.pt <= 9:
                is_insert = True
        except Exception:
            pass
    if not is_insert and run.font.superscript:
        is_insert = not (char_style and char_style in _DOCX_CHAR_STYLE_MAP)

    if is_insert:
        # For Indic-script paragraphs, convert IAST annotation chars → native script.
        # e.g. 'ś' → 'श्', 'g' → 'ग', 'ṅ' → 'ङ्'  (Devanagari)
        # Unknown chars pass through unchanged — so re-imported native annotations
        # (already Devanagari) survive a second import without corruption.
        display = escaped_text
        if para_script == 'devanagari':
            display = ''.join(_IAST_TO_DEV_ANNOTATION.get(c, c) for c in escaped_text)
        # (Telugu / Tamil equivalents can be added here when needed)

        # Superscript insertions (anusvara results, jñ/sv/vy insertions) use <sup>.
        # Non-superscript insertions (inline visarga replacement like ḥ → ś) stay inline.
        if run.font.superscript:
            result = f'<sup><span class="ql-change-style">{display}</span></sup>'
        else:
            result = f'<span class="ql-change-style">{display}</span>'
        return result

    # Apply character style mapping
    if char_style and char_style in _DOCX_CHAR_STYLE_MAP:
        cls, _ = _DOCX_CHAR_STYLE_MAP[char_style]
        # Don't wrap whitespace-only text in style spans — it produces noise
        # (e.g. Svara-style \xa0 spacing runs should stay as plain spaces)
        if escaped_text.strip():
            result = f'<span class="{cls}">{result}</span>'
    elif has_border:
        if border_sz > 6:
            result = f'<span class="ql-holding-long">{result}</span>'
        else:
            result = f'<span class="ql-holding-short">{result}</span>'

    # Bold/italic/underline — only if not already handled by character style
    if not char_style or char_style not in _DOCX_CHAR_STYLE_MAP:
        if run.bold:
            result = f'<strong>{result}</strong>'
        if run.italic:
            result = f'<em>{result}</em>'
        if run.underline:
            result = f'<u>{result}</u>'

    return result


# Character styles that encode holding markers using Unicode combining characters
# (U+034C, U+0342, etc.) rather than w:bdr borders.  The characters themselves
# render as rectangles in browsers; instead, the adjacent base run gets wrapped
# in a CSS holding span and the combining character is dropped.
_DOCX_HOLD_STYLES = frozenset({'hold', '2hold'})
# Character styles that carry reference/footnote markers (f, n, r, m, l …)
# rather than actual text content — skip entirely on import.
_DOCX_SKIP_STYLES = frozenset({
    'phonetic', 'footnote reference', 'endnote reference',
    'footnote text', 'endnote text',
    # Name counters (e.g. '1', '2' in sahasranāma numbering)
    'nāma', 'nama',
    # Empty formatting-only runs
    'pause',
})
# Segments whose raw text is only whitespace or separators should not receive a
# holding — the Hold marker in that position is a look-ahead for the NEXT consonant.
import re as _re
_HOLDING_SEP_RE = _re.compile(r'^[\s\-\xa0\u2002\u2003\u2004]+$')


def _convert_runs_to_html(runs, hanging_indent=False, para_script=None):
    """Convert a list of Word runs to HTML with proper inline formatting.

    Handles Veda Union DOCX character-style encoding:

    Hold / 2Hold runs
      Contain Unicode combining chars (U+034C, U+0342, …) that mark the next
      consonant cluster.  The Hold always precedes its target (look-ahead):
      pending_holding is set and applied to the next non-whitespace base segment.
      The combining character itself is never output (renders as a rectangle).

    Phonetic / footnote reference runs
      Contain annotation markers (f, n, r, m, l …) — skipped entirely.
    """
    tab = '\u2003\u2003' if hanging_indent else ''

    # Flat segment list built before final rendering so look-back is possible.
    # Each entry: {'html': str, 'text': str, 'holdable': bool, 'has_holding': bool}
    segments = []
    pending_holding = None   # look-ahead: holding to apply to next base segment

    for run in runs:
        char_style_name = run.style.name if run.style else 'Default Paragraph Font'
        cs_lower = char_style_name.lower()

        # ── Hold / 2Hold ──────────────────────────────────────────────────────
        # In the Veda Union DOCX encoding, the Hold combining character
        # (U+034C / U+0342) is always placed BEFORE the consonant cluster it
        # marks — it follows the preceding vowel/syllable run, not the target
        # consonant.  Pure look-ahead: set pending_holding and apply it to the
        # next non-whitespace base segment.
        if cs_lower in _DOCX_HOLD_STYLES:
            pending_holding = ('ql-holding-long' if '2' in cs_lower
                               else 'ql-holding-short')
            continue   # never output the combining character itself

        # ── Skip annotation / reference markers ───────────────────────────────
        if cs_lower in _DOCX_SKIP_STYLES:
            continue

        text = run.text
        br_elements = run._element.findall(qn('w:br'))
        has_line_break = len(br_elements) > 0

        if not text and not has_line_break:
            continue

        # "Holdable" = plain base text with no special character style and not
        # superscript.  Svara/accent runs are NOT holdable so the look-back skips
        # them and finds the actual consonant run.
        is_plain_base = (
            cs_lower in ('default paragraph font', '')
            and not run.font.superscript
        )

        def _make_seg(raw_text, formatted_html):
            nonlocal pending_holding
            seg_html = formatted_html
            # Apply look-ahead pending holding to first real (non-separator) base seg
            if (pending_holding and is_plain_base
                    and raw_text.strip()
                    and not _HOLDING_SEP_RE.match(raw_text)):
                seg_html = f'<span class="{pending_holding}">{seg_html}</span>'
                pending_holding = None
            segments.append({
                'html': seg_html,
                'text': raw_text,
                'holdable': is_plain_base,
                'has_holding': seg_html != formatted_html,
            })

        if has_line_break:
            for child in run._element:
                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if tag == 'br':
                    segments.append({'html': f'<br>{tab}', 'text': '\n',
                                     'holdable': False, 'has_holding': False})
                elif tag == 't' and child.text:
                    _make_seg(child.text,
                              _format_run_text(_html_escape(child.text), run, para_script))
            continue

        _make_seg(text, _format_run_text(_html_escape(text), run, para_script))

    return ''.join(seg['html'] for seg in segments)


def _atomic_write_text(path, content, encoding='utf-8'):
    folder = os.path.dirname(path)
    os.makedirs(folder, exist_ok=True)
    tmp = f"{path}.tmp.{os.getpid()}.{threading.get_ident()}"
    with open(tmp, 'w', encoding=encoding) as f:
        f.write(content)
        f.flush()
        os.fsync(f.fileno())
    os.replace(tmp, path)


def _atomic_write_json(path, data):
    _atomic_write_text(path, json.dumps(data, indent=2), encoding='utf-8')


# --- Preferences & Recent Files ---
def load_preferences():
    defaults = {'recents': [], 'recent_files': []}
    if not os.path.exists(PREFS_FILE):
        return defaults
    with PREFS_LOCK:
        try:
            with open(PREFS_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return data if isinstance(data, dict) else defaults
        except Exception:
            return defaults


def save_preferences(prefs):
    with PREFS_LOCK:
        _atomic_write_json(PREFS_FILE, prefs)


def add_to_recent_event(filepath, event='open'):
    """Add or update a file in recent list."""
    if not filepath:
        return
    try:
        prefs = load_preferences()
        # Use 'recents' key (consistent with get_recents and _migrate_recents)
        recents = prefs.get('recents', prefs.get('recent_files', []))
        norm_path = os.path.normpath(os.path.abspath(filepath))
        now = int(time.time())
        
        # Remove existing entry (handle both dict and legacy string items)
        def get_item_path(r):
            if isinstance(r, dict):
                return r.get('path', '')
            elif isinstance(r, str):
                return r
            return ''
        
        # Find existing entry to preserve some fields
        existing = None
        for r in recents:
            if isinstance(r, dict) and os.path.normpath(os.path.abspath(get_item_path(r))) == norm_path:
                existing = r
                break
        
        recents = [r for r in recents if os.path.normpath(os.path.abspath(get_item_path(r))) != norm_path]
        
        # Build new entry with correct field names for get_recents()
        entry = {
            'path': norm_path,
            'last_accessed': now if event == 'open' else (existing.get('last_accessed', 0) if existing else 0),
            'last_saved': now if event == 'save' else (existing.get('last_saved', 0) if existing else 0),
        }
        recents.insert(0, entry)
        recents = recents[:50]  # Keep max 50 entries
        
        # Save to 'recents' key (not 'recent_files')
        prefs['recents'] = recents
        # Remove legacy key if present
        prefs.pop('recent_files', None)
        save_preferences(prefs)
    except Exception as e:
        print(f"Error adding to recent: {e}")


def remove_recent(filepath):
    """Remove a file from recent list."""
    if not filepath:
        return
    try:
        prefs = load_preferences()

        def strip_list(items):
            target_path = os.path.normcase(os.path.normpath(os.path.abspath(filepath)))

            def get_item_path(r):
                if isinstance(r, dict):
                    return r.get('path', '')
                elif isinstance(r, str):
                    return r
                return ''

            return [
                r for r in (items or [])
                if os.path.normcase(os.path.normpath(os.path.abspath(get_item_path(r)))) != target_path
            ]

        prefs['recent_files'] = strip_list(prefs.get('recent_files', []))
        prefs['recents'] = strip_list(prefs.get('recents', []))
        save_preferences(prefs)
    except Exception:
        pass


# --- Flask Application ---
# Disable built-in static file serving by setting static_folder=None
# We handle static files explicitly with the catch-all route at the end
app = Flask(__name__, static_folder=None)
# Allow large audio uploads (~250 MB ≈ 3+ hours of MP3) via /api/audio/editor/state etc.
app.config['MAX_CONTENT_LENGTH'] = 300 * 1024 * 1024

# ─── On-screen keyboard: shared in-memory character queue ───────────────────
import collections as _collections
_keyboard_queue = _collections.deque()
_keyboard_queue_lock = threading.Lock()

# ─── Popup dialog: action queue (dialog → main editor) ──────────────────────
# Popup dialog windows POST events here; the main editor polls and executes them.
_dialog_action_queue = _collections.deque()
_dialog_action_lock = threading.Lock()


@app.route('/')
def index():
    """Serve editor.html with theme injection."""
    try:
        prefs = load_preferences() or {}
        theme_mode = str(prefs.get('theme_mode') or 'system').lower()
        if theme_mode not in ('light', 'dark', 'system'):
            theme_mode = 'system'
        
        editor_path = os.path.join(BASE_DIR, 'editor.html')
        with open(editor_path, 'r', encoding='utf-8') as f:
            html = f.read()
        html = html.replace('__SIKSAMITRA_INITIAL_THEME_MODE__', theme_mode)
        return Response(html, mimetype='text/html')
    except Exception as e:
        return f"Error loading editor: {e}", 500


@app.route('/uploads/<path:filename>')
def serve_uploads(filename):
    """Serve uploaded media files."""
    return send_from_directory(MEDIA_DIR, filename)


# --- File Browser API ---
@app.route('/api/file/browser', methods=['GET'])
def file_browser():
    """Get file browser data including recents and scratch."""
    try:
        prefs = load_preferences()
        recents = prefs.get('recent_files', [])
        
        # Update exists status
        for r in recents:
            r['exists'] = os.path.exists(r.get('path', ''))
        
        # Check for scratch document
        scratch = {'exists': False, 'size': 0, 'modified': 0}
        if os.path.exists(UNTITLED_PATH):
            try:
                stat = os.stat(UNTITLED_PATH)
                scratch = {
                    'exists': True,
                    'size': stat.st_size,
                    'modified': stat.st_mtime
                }
            except Exception:
                pass
        
        return jsonify({
            'recents': recents,
            'scratch': scratch
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/file/read', methods=['POST'])
def read_file_api():
    """Read a file. Handles compressed .smdoc files."""
    data = request.json or {}
    filepath = data.get('path', '').strip()
    
    if not filepath:
        return jsonify({'error': 'Path required'}), 400
    
    try:
        if is_smdoc_file(filepath):
            # Read as binary and decompress if needed
            with open(filepath, 'rb') as f:
                raw_data = f.read()
            content = decompress_smdoc(raw_data)
        else:
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
        add_to_recent_event(filepath, event='open')
        return jsonify({
            'content': content,
            'path': filepath,
            'name': os.path.basename(filepath)
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/file/import-docx', methods=['POST'])
def import_docx_api():
    """Import a .docx file and convert to Quill-compatible HTML."""
    data = request.json or {}
    filepath = data.get('path', '').strip()

    if not filepath:
        return jsonify({'error': 'Path required'}), 400

    if not HAS_DOCX:
        return jsonify({'error': 'python-docx is not installed. Run: pip install python-docx'}), 500

    try:
        # Update loader for live progress
        _update_loader('Importing Word Document', f'Reading {os.path.basename(filepath)}…')
        html_content = convert_docx_to_html(filepath)
        para_count = html_content.count('<p')
        _update_loader('Importing Word Document', f'Converted {para_count} paragraphs — loading into editor…')
        title = os.path.splitext(os.path.basename(filepath))[0]
        return jsonify({
            'content': html_content,
            'path': filepath,
            'name': os.path.basename(filepath),
            'title': title,
            'paragraphs': para_count,
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Failed to import .docx: {str(e)}'}), 500


# ---------------------------------------------------------------------------
# DOCX Export
# ---------------------------------------------------------------------------

# Reverse map: editor CSS class → Word paragraph style name
_EDITOR_CLASS_TO_WORD_STYLE = {
    'ql-doc-title':      'Title',
    'ql-doc-subtitle':   'Subtitle',
    'ql-doc-section':    'Heading 1',
    'ql-doc-subsection': 'Heading 2',
    'ql-doc-translation':'Prijevod',
}

# Script → (font-name, line-height-pt)
_SCRIPT_FONT_MAP = {
    'devanagari': ('Noto Serif Devanagari', 32),
    'telugu':     ('Noto Serif Telugu',     32),
    'tamil':      ('Noto Serif Tamil',      32),
    'kannada':    ('Noto Serif Kannada',    32),
    'malayalam':  ('Noto Serif Malayalam',  32),
}

# Default body font matches "Translit" style in Veda Union docs
_BODY_FONT     = 'Arial'
_BODY_SIZE_PT  = 16
_BODY_LH_PT    = 24   # Translit line-height: lineRule=exact 480 twips
_HEADING_FONT  = 'Arial'


def _docx_set_para_line_height(para, pt):
    """Set exact line height on a python-docx paragraph (in points).
    Uses python-docx's paragraph_format API for reliability."""
    from docx.shared import Pt as _Pt
    from docx.enum.text import WD_LINE_SPACING as _WD_LS
    pf = para.paragraph_format
    pf.line_spacing = _Pt(pt)
    pf.line_spacing_rule = _WD_LS.EXACTLY


def _docx_set_para_spacing(para, before_pt=0, after_pt=0):
    """Set space-before and space-after on a paragraph."""
    from docx.shared import Pt as _Pt
    pf = para.paragraph_format
    pf.space_before = _Pt(before_pt)
    pf.space_after  = _Pt(after_pt)


def _docx_run_set_border(run, size_eighths=4):
    """Add a character border to a run (for holdings)."""
    from docx.oxml.ns import qn as _qn
    import lxml.etree as _etree
    rPr = run._r.get_or_add_rPr()
    bdr = _etree.SubElement(rPr, _qn('w:bdr'))
    bdr.set(_qn('w:val'), 'single')
    bdr.set(_qn('w:sz'), str(size_eighths))
    bdr.set(_qn('w:space'), '0')
    bdr.set(_qn('w:color'), '000000')


class _HtmlToDocxParser:
    """Minimal HTML → python-docx paragraph builder.

    Parses the subset of HTML that śikṣāmitra's Quill editor produces:
      <p [class="..."]>  inline content  </p>
      Inline: <strong>, <em>, <u>, <sup>, <sub>,
              <span class="ql-holding-short|ql-holding-long|ql-change-style|
                          ql-svara-char|ql-doc-translation|ql-translation-style">,
              text nodes, &amp; &lt; &gt; &nbsp;
    """

    def __init__(self, doc):
        self.doc = doc
        self._paragraphs = []   # list of {classes: [...], frags: [...]}
        self._parse_called = False

    # ------------------------------------------------------------------
    def parse(self, html):
        """Parse html string and fill self._paragraphs."""
        import html as _html_mod
        import re

        # Normalise self-closing tags, strip <br> (treated as space)
        html = html.replace('<br>', ' ').replace('<br/>', ' ').replace('<br />', ' ')

        para_pat = re.compile(
            r'<p(?:\s+class=["\']([^"\']*)["\'])?\s*>(.*?)</p>',
            re.DOTALL | re.IGNORECASE
        )

        for m in para_pat.finditer(html):
            classes = (m.group(1) or '').split()
            inner   = m.group(2) or ''
            frags   = self._parse_inline(inner)
            self._paragraphs.append({'classes': classes, 'frags': frags})

    # ------------------------------------------------------------------
    def _parse_inline(self, html):
        """Return list of frag dicts: {text, bold, italic, underline,
        superscript, subscript, classes}."""
        import html as _html_mod
        import re

        frags = []
        # Stack: list of active open tags with their class info
        tag_stack = []

        def current_state():
            bold = italic = underline = superscript = subscript = False
            cls = set()
            for t, tclass in tag_stack:
                if t == 'strong': bold = True
                elif t == 'em':   italic = True
                elif t == 'u':    underline = True
                elif t == 'sup':  superscript = True
                elif t == 'sub':  subscript = True
                elif t == 'span' and tclass:
                    for c in tclass.split():
                        cls.add(c)
            return bold, italic, underline, superscript, subscript, cls

        # Tokenise
        token_pat = re.compile(
            r'(<(?P<close>/)?(?P<tag>strong|em|u|sup|sub|span)(?P<attrs>[^>]*)>)'
            r'|(?P<text>[^<]+)',
            re.DOTALL | re.IGNORECASE
        )

        for m in token_pat.finditer(html):
            if m.group('text') is not None:
                text = _html_mod.unescape(m.group('text'))
                bold, italic, underline, sup, sub, cls = current_state()
                frags.append({
                    'text': text,
                    'bold': bold,
                    'italic': italic,
                    'underline': underline,
                    'superscript': sup,
                    'subscript': sub,
                    'classes': set(cls),
                })
            elif m.group('close'):
                # Closing tag — pop matching from stack
                tag = m.group('tag').lower()
                for i in range(len(tag_stack) - 1, -1, -1):
                    if tag_stack[i][0] == tag:
                        tag_stack.pop(i)
                        break
            else:
                # Opening tag
                tag = m.group('tag').lower()
                attrs = m.group('attrs') or ''
                cls_m = re.search(r'class=["\']([^"\']*)["\']', attrs)
                tclass = cls_m.group(1) if cls_m else ''
                tag_stack.append((tag, tclass))

        return frags

    # ------------------------------------------------------------------
    def build(self, body_font=_BODY_FONT, body_size_pt=_BODY_SIZE_PT,
              body_lh_pt=_BODY_LH_PT):
        """Write parsed paragraphs into self.doc."""
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        for pdata in self._paragraphs:
            classes  = pdata['classes']
            frags    = pdata['frags']

            # Determine Word paragraph style
            word_style = 'Normal'
            for cls in classes:
                if cls in _EDITOR_CLASS_TO_WORD_STYLE:
                    word_style = _EDITOR_CLASS_TO_WORD_STYLE[cls]
                    break

            # Determine script (for font selection)
            script_cls = next((c.replace('ql-script-', '')
                               for c in classes if c.startswith('ql-script-')), None)
            is_translation = ('ql-doc-translation' in classes or
                              'ql-translation-style' in classes)

            # Choose font and line height for this paragraph
            if script_cls and script_cls in _SCRIPT_FONT_MAP:
                para_font, para_lh = _SCRIPT_FONT_MAP[script_cls]
            elif is_translation:
                para_font, para_lh = ('Times New Roman', 18)
            elif word_style.startswith('Heading'):
                para_font, para_lh = (_HEADING_FONT, 32)
            elif word_style == 'Title':
                para_font, para_lh = (_HEADING_FONT, 40)
            else:
                para_font, para_lh = (body_font, body_lh_pt)

            # Add paragraph (use Normal to avoid inheriting heading numbering etc.)
            try:
                para = self.doc.add_paragraph(style=word_style)
            except Exception:
                para = self.doc.add_paragraph(style='Normal')

            _docx_set_para_line_height(para, para_lh)
            _docx_set_para_spacing(para, before_pt=0, after_pt=0)

            # Add runs
            for frag in frags:
                if not frag['text']:
                    continue
                run = para.add_run(frag['text'])

                # Font
                run.font.name = para_font
                if script_cls and script_cls in _SCRIPT_FONT_MAP:
                    run.font.name = _SCRIPT_FONT_MAP[script_cls][0]

                # Size
                frag_size = body_size_pt
                if is_translation:
                    frag_size = 11
                elif frag['superscript'] or 'ql-change-style' in frag['classes']:
                    frag_size = 8
                elif word_style == 'Heading 1':
                    frag_size = 24
                elif word_style == 'Heading 2':
                    frag_size = 22
                elif word_style == 'Title':
                    frag_size = 24
                run.font.size = Pt(frag_size)

                # Basic formatting
                fclasses = frag['classes']
                run.bold      = frag['bold']
                run.italic    = (frag['italic'] or is_translation or
                                 'ql-change-style' in fclasses or
                                 'ql-translation-style' in fclasses)
                run.underline = frag['underline']
                run.font.superscript = frag['superscript']
                run.font.subscript   = frag['subscript']

                # Holdings → character border
                if 'ql-holding-long' in fclasses:
                    _docx_run_set_border(run, size_eighths=8)
                elif 'ql-holding-short' in fclasses:
                    _docx_run_set_border(run, size_eighths=4)

                # Svara accent → red bold
                if 'ql-svara-char' in fclasses or 'ql-svara' in fclasses:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xCC, 0x1B, 0x1B)

                # Change style → blue italic
                if 'ql-change-style' in fclasses:
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

                # Translation → gray
                if is_translation or 'ql-translation-style' in fclasses:
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def convert_html_to_docx(html_content, title=''):
    """Convert śikṣāmitra editor HTML to a .docx file (as bytes).

    Produces a document matching the Veda Union layout:
      - A4 page (210mm × 297mm)
      - Margins: left 25mm, right 9mm, top 20mm, bottom 15mm
      - Body text: Arial 16pt, exact 24pt line height (Translit style)
      - Indic script paragraphs: Noto Serif font, 32pt line height
      - Headings: Arial, sizes per _EDITOR_CLASS_TO_WORD_STYLE mapping
      - Translations: Times New Roman 11pt italic gray
    """
    if not HAS_DOCX:
        raise ImportError("python-docx is required for DOCX export")

    from docx.shared import Mm, Pt, RGBColor
    from io import BytesIO

    doc = DocxDocument()

    # --- Page layout: A4, Veda Union margins ---
    section = doc.sections[0]
    section.page_width  = Mm(210)
    section.page_height = Mm(297)
    section.left_margin   = Mm(25)
    section.right_margin  = Mm(9)
    section.top_margin    = Mm(20)
    section.bottom_margin = Mm(15)

    # --- Define / override styles ---
    styles = doc.styles

    def _ensure_style(name, base='Normal'):
        try:
            return styles[name]
        except KeyError:
            from docx.enum.style import WD_STYLE_TYPE
            return styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH, builtin=False)

    # Translit (body mantra text) style
    translit = _ensure_style('Translit', base='Normal')
    translit.font.name = _BODY_FONT
    translit.font.size = Pt(_BODY_SIZE_PT)
    pf = translit.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    pf.line_spacing = Pt(_BODY_LH_PT)
    from docx.enum.text import WD_LINE_SPACING as _WD_LS
    pf.line_spacing_rule = _WD_LS.EXACTLY

    # Prijevod (translation) style — only create if absent
    prijevod = _ensure_style('Prijevod', base='Normal')
    prijevod.font.name = 'Times New Roman'
    prijevod.font.size = Pt(11)
    prijevod.font.italic = True
    prijevod.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # --- Parse and write paragraphs ---
    parser = _HtmlToDocxParser(doc)
    parser.parse(html_content)
    parser.build(body_font=_BODY_FONT, body_size_pt=_BODY_SIZE_PT,
                 body_lh_pt=_BODY_LH_PT)

    # --- Remove the blank paragraph that python-docx adds by default ---
    # (the first paragraph in a new document is always empty)
    if doc.paragraphs and not doc.paragraphs[0].text:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


@app.route('/api/file/export-docx', methods=['POST'])
def export_docx_api():
    """Export editor HTML content to a .docx file.

    Request JSON: { "content": "<html>", "path": "/save/to.docx", "title": "..." }
    Saves the file server-side and returns { "path": ... }.
    """
    data = request.json or {}
    content  = data.get('content', '')
    filepath = (data.get('path', '') or '').strip()
    title    = data.get('title', '')

    if not content:
        return jsonify({'error': 'No content provided'}), 400
    if not filepath:
        return jsonify({'error': 'Path required'}), 400
    if not HAS_DOCX:
        return jsonify({'error': 'python-docx is not installed. Run: pip install python-docx'}), 500

    try:
        docx_bytes = convert_html_to_docx(content, title=title)
        os.makedirs(os.path.dirname(os.path.abspath(filepath)), exist_ok=True)
        with open(filepath, 'wb') as f:
            f.write(docx_bytes)
        return jsonify({'path': filepath, 'size': len(docx_bytes)})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'DOCX export failed: {e}'}), 500


@app.route('/api/file/save', methods=['POST'])
def save_file_api():
    """Save a file. Compresses .smdoc files for efficiency."""
    data = request.json or {}
    filepath = data.get('path', '').strip()
    content = data.get('content', '')
    
    if not filepath:
        return jsonify({'error': 'Path required'}), 400
    
    try:
        folder = os.path.dirname(filepath)
        if folder:
            os.makedirs(folder, exist_ok=True)
        
        if is_smdoc_file(filepath):
            # Compress .smdoc files with maximum compression
            compressed_data = compress_smdoc(content)
            with open(filepath, 'wb') as f:
                f.write(compressed_data)
        else:
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content)
        add_to_recent_event(filepath, event='save')
        return jsonify({'status': 'saved', 'path': filepath})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/file/delete', methods=['POST'])
def delete_file_api():
    """Delete a file."""
    data = request.json or {}
    filepath = data.get('path', '').strip()
    
    if not filepath or not os.path.exists(filepath):
        return jsonify({'error': 'File not found'}), 404
    
    try:
        if os.path.isdir(filepath):
            shutil.rmtree(filepath)
        else:
            os.remove(filepath)
        remove_recent(filepath)
        return jsonify({'status': 'deleted'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# --- Unified Library Endpoint (for File tab) ---
def _normalize_recent_path(path):
    try:
        return os.path.normcase(os.path.abspath(path))
    except Exception:
        return str(path).lower()


def _migrate_recents(prefs):
    """Migrate legacy recents format."""
    if 'recents' not in prefs and 'recent_files' in prefs:
        prefs['recents'] = prefs.pop('recent_files')
    return prefs


def get_recents(prefs):
    """Get recents list with proper format."""
    recents = prefs.get('recents', prefs.get('recent_files', []))
    if not isinstance(recents, list):
        return []
    cleaned = []
    for r in recents:
        # Handle dict items
        if isinstance(r, dict):
            p = r.get('path')
            if isinstance(p, str) and p.strip():
                cleaned.append({
                    'path': p,
                    'last_accessed': int(r.get('last_accessed') or 0),
                    'last_saved': int(r.get('last_saved') or 0),
                })
        # Handle legacy string items
        elif isinstance(r, str) and r.strip():
            cleaned.append({
                'path': r,
                'last_accessed': 0,
                'last_saved': 0,
            })
    return cleaned


def build_recent_item(path, meta=None):
    """Build a recent item with metadata."""
    meta = meta or {}
    abs_path = os.path.abspath(path)
    exists = os.path.exists(abs_path)
    modified = 0
    size = 0
    if exists:
        try:
            stat = os.stat(abs_path)
            modified = stat.st_mtime
            size = stat.st_size
        except Exception:
            modified = 0
            size = 0

    try:
        is_internal = _normalize_recent_path(abs_path).startswith(_normalize_recent_path(LIBRARY_DIR) + os.sep)
    except Exception:
        is_internal = False

    return {
        'path': abs_path,
        'name': os.path.basename(abs_path),
        'dir': os.path.dirname(abs_path),
        'modified': modified,
        'size': size,
        'exists': bool(exists),
        'is_internal': bool(is_internal),
        'last_accessed': int(meta.get('last_accessed') or 0),
        'last_saved': int(meta.get('last_saved') or 0),
    }


@app.route('/api/recents/remove', methods=['POST'])
def api_remove_recent():
    """Remove a file from recents list."""
    data = request.json or {}
    path = data.get('path')
    if not path or not isinstance(path, str):
        return jsonify({'error': 'path required'}), 400
    try:
        remove_recent(path)
        return jsonify({'status': 'removed'})
    except Exception as exc:
        return jsonify({'error': str(exc)}), 500


@app.route('/api/library', methods=['GET'])
def get_library():
    """Unified payload for the File tab (session + recents)."""
    prefs = load_preferences()
    prefs = _migrate_recents(prefs)
    recents_meta = get_recents(prefs)
    # Sort by most recent activity: prefer last_saved (edit), fall back to last_accessed (open)
    recents_meta.sort(key=lambda r: max(r.get('last_saved') or 0, r.get('last_accessed') or 0), reverse=True)
    recents = [build_recent_item(r['path'], r) for r in recents_meta[:20]]

    session = None
    if os.path.exists(SESSION_STATE_PATH):
        try:
            with open(SESSION_STATE_PATH, 'r', encoding='utf-8') as f:
                session = json.load(f)
        except Exception:
            session = None

    scratch = {
        'exists': False,
        'modified': 0,
        'size': 0,
        'path': UNTITLED_PATH,
    }
    try:
        if os.path.exists(UNTITLED_PATH):
            scratch['exists'] = True
            scratch['modified'] = os.path.getmtime(UNTITLED_PATH)
            scratch['size'] = os.path.getsize(UNTITLED_PATH)
    except Exception:
        pass

    return jsonify({'session': session, 'scratch': scratch, 'recents': recents})


# --- Library API ---
def _build_library_tree(path, relative_base=''):
    """Recursively build library tree."""
    items = []
    try:
        entries = sorted(os.listdir(path), key=lambda x: (not os.path.isdir(os.path.join(path, x)), x.lower()))
        for name in entries:
            full_path = os.path.join(path, name)
            rel_path = os.path.join(relative_base, name).replace('\\', '/')
            item = {
                'name': name,
                'path': rel_path,
                'full_path': full_path,
                'is_folder': os.path.isdir(full_path),
                'modified': 0,
                'size': 0
            }
            if item['is_folder']:
                item['children'] = _build_library_tree(full_path, rel_path)
            else:
                try:
                    stat = os.stat(full_path)
                    item['modified'] = stat.st_mtime
                    item['size'] = stat.st_size
                except Exception:
                    pass
            items.append(item)
    except Exception:
        pass
    return items


@app.route('/api/library/browse', methods=['GET'])
def browse_library():
    """Get library folder tree."""
    try:
        tree = _build_library_tree(LIBRARY_DIR)
        return jsonify({'root': LIBRARY_DIR, 'items': tree})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/library/folder', methods=['POST'])
def create_library_folder():
    """Create folder in library."""
    data = request.json or {}
    rel_path = data.get('path', '').strip()
    name = data.get('name', '').strip()
    
    if not name:
        return jsonify({'error': 'Folder name required'}), 400
    
    name = re.sub(r'[<>:"/\\|?*]', '_', name)
    parent = os.path.join(LIBRARY_DIR, rel_path) if rel_path else LIBRARY_DIR
    new_folder = os.path.join(parent, name)
    
    if os.path.exists(new_folder):
        return jsonify({'error': 'Folder already exists'}), 400
    
    try:
        os.makedirs(new_folder, exist_ok=True)
        return jsonify({'status': 'created', 'path': os.path.relpath(new_folder, LIBRARY_DIR).replace('\\', '/')})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/library/rename', methods=['POST'])
def rename_library_item():
    """Rename file/folder in library."""
    data = request.json or {}
    rel_path = data.get('path', '').strip()
    new_name = data.get('name', '').strip()
    
    if not rel_path or not new_name:
        return jsonify({'error': 'Path and name required'}), 400
    
    old_full = os.path.join(LIBRARY_DIR, rel_path)
    if not os.path.exists(old_full):
        return jsonify({'error': 'Item not found'}), 404
    
    new_name = re.sub(r'[<>:"/\\|?*]', '_', new_name)
    new_full = os.path.join(os.path.dirname(old_full), new_name)
    
    if os.path.exists(new_full):
        return jsonify({'error': 'Name already exists'}), 400
    
    try:
        os.rename(old_full, new_full)
        return jsonify({'status': 'renamed', 'new_path': os.path.relpath(new_full, LIBRARY_DIR).replace('\\', '/')})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/library/move', methods=['POST'])
def move_library_item():
    """Move file/folder to a different folder in library."""
    data = request.json or {}
    rel_path = data.get('path', '').strip()
    target_folder = data.get('target', '').strip()
    
    if not rel_path:
        return jsonify({'error': 'Path required'}), 400
    
    source_full = os.path.join(LIBRARY_DIR, rel_path)
    if not os.path.exists(source_full):
        return jsonify({'error': 'Item not found'}), 404
    
    # Target can be empty (root) or a subfolder
    if target_folder:
        target_dir = os.path.join(LIBRARY_DIR, target_folder)
    else:
        target_dir = LIBRARY_DIR
    
    if not os.path.isdir(target_dir):
        return jsonify({'error': 'Target folder not found'}), 404
    
    # Security check - both must be inside LIBRARY_DIR
    if not os.path.abspath(source_full).startswith(os.path.abspath(LIBRARY_DIR)):
        return jsonify({'error': 'Invalid source path'}), 400
    if not os.path.abspath(target_dir).startswith(os.path.abspath(LIBRARY_DIR)):
        return jsonify({'error': 'Invalid target path'}), 400
    
    # Can't move a folder into itself or its children
    if os.path.isdir(source_full):
        if os.path.abspath(target_dir).startswith(os.path.abspath(source_full)):
            return jsonify({'error': 'Cannot move folder into itself'}), 400
    
    item_name = os.path.basename(source_full)
    new_full = os.path.join(target_dir, item_name)
    
    if os.path.exists(new_full):
        return jsonify({'error': 'Item already exists in target folder'}), 400
    
    try:
        shutil.move(source_full, new_full)
        new_rel = os.path.relpath(new_full, LIBRARY_DIR).replace('\\', '/')
        return jsonify({'status': 'moved', 'new_path': new_rel})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/library/delete', methods=['POST'])
def delete_library_item():
    """Delete file/folder in library."""
    data = request.json or {}
    rel_path = data.get('path', '').strip()
    
    if not rel_path:
        return jsonify({'error': 'Path required'}), 400
    
    full_path = os.path.join(LIBRARY_DIR, rel_path)
    
    if not os.path.exists(full_path):
        return jsonify({'error': 'Item not found'}), 404
    
    # Security check
    if not os.path.abspath(full_path).startswith(os.path.abspath(LIBRARY_DIR)):
        return jsonify({'error': 'Invalid path'}), 400
    
    try:
        if os.path.isdir(full_path):
            shutil.rmtree(full_path)
        else:
            os.remove(full_path)
        return jsonify({'status': 'deleted'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/library/save', methods=['POST'])
def save_to_library():
    """Save document to library. Compresses .smdoc files."""
    data = request.json or {}
    rel_path = data.get('path', '').strip()
    filename = data.get('filename', 'document.html').strip()
    content = data.get('content', '')
    
    if not filename:
        filename = 'document.html'
    # Don't auto-add .html extension if it's already .smdoc
    if not filename.lower().endswith('.html') and not filename.lower().endswith('.smdoc'):
        filename += '.html'
    
    parent = os.path.join(LIBRARY_DIR, rel_path) if rel_path else LIBRARY_DIR
    full_path = os.path.join(parent, filename)
    
    try:
        os.makedirs(parent, exist_ok=True)
        if is_smdoc_file(full_path):
            # Compress .smdoc files
            compressed_data = compress_smdoc(content)
            with open(full_path, 'wb') as f:
                f.write(compressed_data)
        else:
            with open(full_path, 'w', encoding='utf-8') as f:
                f.write(content)
        add_to_recent_event(full_path, event='save')
        return jsonify({
            'status': 'saved',
            'path': os.path.relpath(full_path, LIBRARY_DIR).replace('\\', '/'),
            'full_path': full_path
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/library/read', methods=['POST'])
def read_library_file():
    """Read file from library. Handles compressed .smdoc files."""
    data = request.json or {}
    rel_path = data.get('path', '').strip()
    
    if not rel_path:
        return jsonify({'error': 'Path required'}), 400
    
    full_path = os.path.join(LIBRARY_DIR, rel_path)
    
    if not os.path.exists(full_path):
        return jsonify({'error': 'File not found'}), 404
    
    if os.path.isdir(full_path):
        return jsonify({'error': 'Cannot read folder'}), 400
    
    try:
        if is_smdoc_file(full_path):
            # Read and decompress .smdoc files
            with open(full_path, 'rb') as f:
                raw_data = f.read()
            content = decompress_smdoc(raw_data)
        else:
            with open(full_path, 'r', encoding='utf-8') as f:
                content = f.read()
        add_to_recent_event(full_path, event='open')
        return jsonify({'content': content, 'path': rel_path, 'full_path': full_path})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/library/import', methods=['POST'])
def import_to_library():
    """Import external file to library."""
    data = request.json or {}
    source_path = data.get('source_path', '').strip()
    dest_folder = data.get('dest_folder', '').strip()
    as_link = data.get('as_link', False)
    
    if not source_path or not os.path.exists(source_path):
        return jsonify({'error': 'Source file not found'}), 404
    
    filename = os.path.basename(source_path)
    dest_parent = os.path.join(LIBRARY_DIR, dest_folder) if dest_folder else LIBRARY_DIR
    dest_path = os.path.join(dest_parent, filename)
    
    try:
        os.makedirs(dest_parent, exist_ok=True)
        
        if as_link:
            # Create a .link file with the path
            link_path = dest_path + '.link'
            with open(link_path, 'w', encoding='utf-8') as f:
                json.dump({'source': source_path, 'name': filename}, f)
            return jsonify({
                'status': 'linked',
                'path': os.path.relpath(link_path, LIBRARY_DIR).replace('\\', '/')
            })
        else:
            # Copy the file
            shutil.copy2(source_path, dest_path)
            add_to_recent_event(dest_path, event='import')
            return jsonify({
                'status': 'imported',
                'path': os.path.relpath(dest_path, LIBRARY_DIR).replace('\\', '/'),
                'full_path': dest_path
            })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# --- Scratch/Session API ---
@app.route('/api/scratch/save', methods=['POST'])
def save_scratch():
    """Save scratch document."""
    data = request.json or {}
    content = data.get('content', '')
    try:
        with open(UNTITLED_PATH, 'w', encoding='utf-8') as f:
            f.write(content)
        return jsonify({'status': 'saved'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/scratch/load', methods=['GET'])
def load_scratch():
    """Load scratch document."""
    try:
        if not os.path.exists(UNTITLED_PATH):
            return jsonify({'content': '', 'exists': False})
        with open(UNTITLED_PATH, 'r', encoding='utf-8') as f:
            content = f.read()
        return jsonify({'content': content, 'exists': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/scratch/clear', methods=['POST'])
def clear_scratch():
    """Clear scratch document."""
    try:
        if os.path.exists(UNTITLED_PATH):
            os.remove(UNTITLED_PATH)
        return jsonify({'status': 'cleared'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# --- Cache API (unified endpoints for JS compatibility) ---
@app.route('/api/cache/untitled', methods=['GET', 'POST', 'DELETE'])
def cache_untitled():
    """Handle untitled/scratch document cache."""
    if request.method == 'GET':
        try:
            if not os.path.exists(UNTITLED_PATH):
                return jsonify({'content': '', 'exists': False})
            with open(UNTITLED_PATH, 'r', encoding='utf-8') as f:
                content = f.read()
            # Get file size for the scratch info
            size = os.path.getsize(UNTITLED_PATH) if os.path.exists(UNTITLED_PATH) else 0
            modified = os.path.getmtime(UNTITLED_PATH) if os.path.exists(UNTITLED_PATH) else 0
            return jsonify({
                'content': content, 
                'exists': True,
                'size': size,
                'modified': modified
            })
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    elif request.method == 'POST':
        data = request.json or {}
        content = data.get('content', '')
        try:
            with open(UNTITLED_PATH, 'w', encoding='utf-8') as f:
                f.write(content)
            return jsonify({'status': 'saved'})
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    else:  # DELETE
        try:
            if os.path.exists(UNTITLED_PATH):
                os.remove(UNTITLED_PATH)
            return jsonify({'status': 'cleared'})
        except Exception as e:
            return jsonify({'error': str(e)}), 500


@app.route('/api/cache/session', methods=['GET', 'POST'])
def cache_session():
    """Handle session state cache."""
    if request.method == 'GET':
        try:
            if os.path.exists(SESSION_STATE_PATH):
                with open(SESSION_STATE_PATH, 'r', encoding='utf-8') as f:
                    return jsonify(json.load(f))
            return jsonify({})
        except Exception:
            return jsonify({})
    else:
        try:
            data = request.json or {}
            _atomic_write_json(SESSION_STATE_PATH, data)
            return jsonify({'status': 'saved'})
        except Exception as e:
            return jsonify({'error': str(e)}), 500


@app.route('/api/session/state', methods=['GET', 'POST'])
def session_state():
    """Get or save session state."""
    if request.method == 'GET':
        try:
            if os.path.exists(SESSION_STATE_PATH):
                with open(SESSION_STATE_PATH, 'r', encoding='utf-8') as f:
                    return jsonify(json.load(f))
            return jsonify({})
        except Exception:
            return jsonify({})
    else:
        try:
            data = request.json or {}
            _atomic_write_json(SESSION_STATE_PATH, data)
            return jsonify({'status': 'saved'})
        except Exception as e:
            return jsonify({'error': str(e)}), 500


# --- Preferences API ---
@app.route('/api/preferences', methods=['GET', 'POST'])
def preferences_api():
    """Get or save preferences."""
    if request.method == 'GET':
        return jsonify(load_preferences())
    else:
        try:
            data = request.json or {}
            save_preferences(data)
            return jsonify({'status': 'saved'})
        except Exception as e:
            return jsonify({'error': str(e)}), 500


@app.route('/api/preferences/<key>', methods=['GET', 'PUT'])
def preference_key(key):
    """Get or set single preference."""
    prefs = load_preferences()
    if request.method == 'GET':
        return jsonify({'value': prefs.get(key)})
    else:
        data = request.json or {}
        prefs[key] = data.get('value')
        save_preferences(prefs)
        return jsonify({'status': 'saved'})


# --- Media API ---
@app.route('/api/media/upload', methods=['POST'])
def upload_media():
    """Upload media file."""
    if 'file' not in request.files:
        return jsonify({'error': 'No file'}), 400
    
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': 'No filename'}), 400
    
    filename = re.sub(r'[^\w\-_\.]', '_', file.filename)
    filepath = os.path.join(MEDIA_DIR, filename)
    
    # Handle duplicates
    base, ext = os.path.splitext(filename)
    counter = 1
    while os.path.exists(filepath):
        filename = f"{base}_{counter}{ext}"
        filepath = os.path.join(MEDIA_DIR, filename)
        counter += 1
    
    file.save(filepath)
    return jsonify({
        'url': f'/uploads/{filename}',
        'filename': filename,
        'path': filepath
    })


@app.route('/api/media/list', methods=['GET'])
def list_media():
    """List media files."""
    files = []
    for filename in os.listdir(MEDIA_DIR):
        filepath = os.path.join(MEDIA_DIR, filename)
        if os.path.isfile(filepath):
            files.append({
                'name': filename,
                'size': os.path.getsize(filepath),
                'url': f'/uploads/{filename}',
                'modified': os.path.getmtime(filepath)
            })
    files.sort(key=lambda x: x['modified'], reverse=True)
    return jsonify(files)


# --- Sanskrit Documents Proxy ---

SANSKRITDOCS_BASE = 'https://sanskritdocuments.org'
SANSKRITDOCS_CATEGORIES = [
    ('doc_veda', 'Veda'),
    ('doc_upanishhat', 'Upaniṣad'),
    ('doc_vishhnu', 'Viṣṇu'),
    ('doc_shiva', 'Śiva'),
    ('doc_devii', 'Devī'),
    ('doc_ganesha', 'Gaṇeśa'),
    ('doc_giitaa', 'Gītā'),
    ('doc_raama', 'Rāma'),
    ('doc_hanumaana', 'Hanumān'),
    ('doc_surya', 'Sūrya'),
    ('doc_gurudev', 'Guru'),
    ('doc_subhaashita', 'Subhāṣita'),
    ('doc_z_misc_shlokas', 'Ślokas'),
    ('doc_z_misc_general', 'General'),
    ('doc_z_misc_major_works', 'Major Works'),
    ('doc_z_misc_navagraha', 'Navagraha'),
    ('doc_z_misc_purana', 'Purāṇa'),
    ('doc_z_misc_articles', 'Articles'),
]

_sanskritdocs_index_cache = {}  # category -> [(filename, devanagari_title, romanized_title)]


def _fetch_url(url, timeout=15):
    """Fetch a URL and return decoded text."""
    req = urllib.request.Request(url, headers={
        'User-Agent': 'SiksamitraEditor/1.0',
        'Accept': 'text/html,application/xhtml+xml,*/*',
    })
    with urllib.request.urlopen(req, timeout=timeout) as resp:
        return resp.read().decode('utf-8', errors='replace')


def _parse_index_page(html, category):
    """Parse a category index page and extract document entries.

    The index pages use a structure like:
      <a href="/doc_veda/puruSukta.html" title="...">
        <em lang="sa">पुरुषसूक्तम्</em></a> | <em itemprop="name">Purushasukta</em>
    We need to extract:
      - filename from the href
      - Devanagari title from inside the <a> tag (may be in <em> child)
      - Romanized title from <em itemprop="name"> or the title attribute
    """
    entries = []
    seen = set()

    # Match <a> tags linking to .html in this category, capturing their full inner HTML
    link_pattern = re.compile(
        r'<a[^>]+href=["\'](?:[^"\']*?' + re.escape(category) + r'/)?' +
        r'([^"\']+?)\.html["\'][^>]*>(.*?)</a>' +
        r'(?:\s*\|?\s*(?:<em[^>]*>([^<]*)</em>))?',
        re.IGNORECASE | re.DOTALL
    )

    for match in link_pattern.finditer(html):
        filename = match.group(1).strip()
        inner_html = match.group(2).strip()
        itemprop_name = (match.group(3) or '').strip()

        if not filename or filename.startswith('http') or '/' in filename:
            continue
        if filename in seen:
            continue
        if filename in ('index', 'home', 'favicon'):
            continue

        # Extract text from inner HTML (strip tags)
        inner_text = re.sub(r'<[^>]+>', '', inner_html).strip()

        # Skip empty icon-only links (just whitespace)
        if not inner_text and not itemprop_name:
            continue

        # Build the display title: prefer Devanagari, fall back to romanized
        title = inner_text if inner_text else itemprop_name
        romanized = itemprop_name if itemprop_name else filename

        seen.add(filename)
        entries.append((filename, title, romanized))

    return entries


@app.route('/api/shlokam/search', methods=['GET'])
def shlokam_search():
    """Proxy the Shlokam.org WordPress search API to bypass CORS from dialog windows.

    Shlokam may return different payload shapes for different queries. We normalize
    to the shape the shloka dialog expects: [{ id, title, url, subtype }, ...].
    """
    query = (request.args.get('q') or '').strip()
    if not query:
        return jsonify([])
    try:
        import json as _json
        url = f'https://shlokam.org/wp-json/wp/v2/search?search={urllib.parse.quote(query)}&per_page=20'
        raw = _fetch_url(url, timeout=15)
        try:
            data = _json.loads(raw)
        except Exception:
            return jsonify({'error': 'Invalid JSON from shlokam.org'}), 502
        if not isinstance(data, list):
            return jsonify([])
        out = []
        import html as _html_mod
        for item in data:
            if not isinstance(item, dict):
                continue
            link = item.get('url') or item.get('link')
            if not link:
                continue
            out.append({
                'id': item.get('id'),
                'title': _html_mod.unescape(item.get('title') or item.get('title_plain') or ''),
                'url': link,
                'subtype': item.get('subtype') or item.get('type') or 'post',
            })
        return jsonify(out)
    except Exception as e:
        return jsonify({'error': f'Shlokam search failed: {e}'}), 502


@app.route('/api/shlokam/details/<collection>/<int:post_id>', methods=['GET'])
def shlokam_details(collection, post_id):
    """Proxy a Shlokam.org post/page and extract the individual verses."""
    collection = re.sub(r'[^a-z_]', '', (collection or 'posts').lower()) or 'posts'
    try:
        import json as _json
        url = (f'https://shlokam.org/wp-json/wp/v2/{collection}/{post_id}'
               f'?_fields=content.rendered,title.rendered,link')
        raw = _fetch_url(url, timeout=20)
        data = _json.loads(raw)
        html_content = ''
        if isinstance(data, dict):
            content = data.get('content') or {}
            if isinstance(content, dict):
                html_content = content.get('rendered') or ''
        # Extract verses using regex (basic but robust)
        def _extract(cls):
            import re as _re
            pat = _re.compile(r'<div[^>]*class="[^"]*' + _re.escape(cls) + r'[^"]*"[^>]*>(.*?)</div>',
                              _re.DOTALL | _re.IGNORECASE)
            blocks = []
            for m in pat.findall(html_content or ''):
                # Strip tags and normalize whitespace
                txt = _re.sub(r'<[^>]+>', '', m)
                import html as _html
                txt = _html.unescape(txt)
                blocks.append(txt.strip())
            return blocks
        dev = _extract('verse_sanskrit')
        translit = _extract('verse_trans')
        meaning = _extract('verse_meaning')
        n = max(len(dev), len(translit), len(meaning))
        shlokas = []
        for i in range(n):
            d = dev[i] if i < len(dev) else ''
            t = translit[i] if i < len(translit) else ''
            m = meaning[i] if i < len(meaning) else ''
            if d or t or m:
                shlokas.append({'devanagari': d, 'transliteration': t, 'translation': m})
        return jsonify({'shlokas': shlokas})
    except Exception as e:
        return jsonify({'error': f'Shlokam details failed: {e}'}), 502


@app.route('/api/sanskritdocs/categories', methods=['GET'])
def sanskritdocs_categories():
    """Return the list of document categories."""
    return jsonify([{'id': cid, 'name': cname} for cid, cname in SANSKRITDOCS_CATEGORIES])


@app.route('/api/sanskritdocs/index/<category>', methods=['GET'])
def sanskritdocs_index(category):
    """Fetch and parse the index page for a category, with caching."""
    # Validate category
    valid_ids = {c[0] for c in SANSKRITDOCS_CATEGORIES}
    if category not in valid_ids:
        return jsonify({'error': 'Invalid category'}), 400

    if category in _sanskritdocs_index_cache:
        entries = _sanskritdocs_index_cache[category]
    else:
        try:
            html = _fetch_url(f'{SANSKRITDOCS_BASE}/{category}/')
            entries = _parse_index_page(html, category)
            _sanskritdocs_index_cache[category] = entries
        except Exception as e:
            return jsonify({'error': f'Failed to fetch index: {e}'}), 502

    # Search filter
    query = request.args.get('q', '').strip().lower()
    if query:
        results = []
        for filename, title, romanized in entries:
            search_text = f'{title} {romanized} {filename}'.lower()
            if query in search_text:
                results.append((filename, title, romanized))
                continue
            # Fuzzy: check if all words in query appear somewhere
            words = query.split()
            if all(w in search_text for w in words):
                results.append((filename, title, romanized))
        entries = results

    return jsonify([
        {'filename': f, 'title': t, 'romanized': r, 'category': category}
        for f, t, r in entries
    ])


@app.route('/api/sanskritdocs/fetch/<category>/<filename>', methods=['GET'])
def sanskritdocs_fetch(category, filename):
    """Fetch a specific document's HTML and extract the Sanskrit text."""
    valid_ids = {c[0] for c in SANSKRITDOCS_CATEGORIES}
    if category not in valid_ids:
        return jsonify({'error': 'Invalid category'}), 400

    # Sanitize filename
    filename = re.sub(r'[^a-zA-Z0-9_\-]', '', filename)
    if not filename:
        return jsonify({'error': 'Invalid filename'}), 400

    try:
        html = _fetch_url(f'{SANSKRITDOCS_BASE}/{category}/{filename}.html')
    except Exception as e:
        return jsonify({'error': f'Failed to fetch document: {e}'}), 502

    # Extract Sanskrit text from HTML
    import html as html_mod
    text = html
    text = re.sub(r'<script[^>]*>.*?</script>', '', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<br\s*/?\s*>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'</p>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'</div>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'<[^>]+>', '', text)
    text = html_mod.unescape(text)

    # Filter to lines with substantial Devanagari content
    # This removes navigation ("Home", "ITX", "PDF") while keeping actual text
    content_lines = []
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
        dev_chars = len(re.findall(r'[\u0900-\u097F]', line))
        total_chars = len(line.replace(' ', ''))
        if dev_chars >= 3 and (dev_chars / max(total_chars, 1)) > 0.3:
            content_lines.append(line)

    # Remove footer metadata lines
    footer_keywords = [
        'encoded', 'proofread', 'send corrections',
        'sanskritdocuments.org', 'last updated', 'latest update',
        '% text title', '% file name'
    ]
    while content_lines:
        if any(kw in content_lines[-1].lower() for kw in footer_keywords):
            content_lines.pop()
        else:
            break

    # Separate preamble (title, ṛṣi/chandas metadata) from mantra text.
    # Preamble lines typically contain commas, latin chars, or lack Vedic
    # svara marks (U+0951, U+0952, U+1CDA). Actual mantras have svara marks.
    mantra_lines = []
    preamble_lines = []
    has_svara = lambda l: bool(re.search(r'[\u0951\u0952\u1CDA]', l))

    # Find where mantra text begins (first line with svara marks)
    mantra_started = False
    for line in content_lines:
        if not mantra_started:
            if has_svara(line):
                mantra_started = True
                mantra_lines.append(line)
            else:
                preamble_lines.append(line)
        else:
            mantra_lines.append(line)

    # If no svara marks found at all (non-Vedic text), use all lines
    if not mantra_lines:
        mantra_lines = content_lines
        preamble_lines = []

    full_text = '\n'.join(mantra_lines)
    preamble_text = '\n'.join(preamble_lines)

    # Extract title from the page
    title_match = re.search(r'<title[^>]*>([^<]+)</title>', html, re.IGNORECASE)
    title = title_match.group(1).strip() if title_match else filename

    return jsonify({
        'text': full_text,
        'preamble': preamble_text,
        'title': title,
        'filename': filename,
        'category': category,
        'url': f'{SANSKRITDOCS_BASE}/{category}/{filename}.html'
    })


# --- Flask Server Thread ---
class FlaskServerThread(threading.Thread):
    def __init__(self, flask_app, host='127.0.0.1', port=0):
        super().__init__(daemon=True)
        self._server = make_server(host, port, flask_app, threaded=True)
        self.host = host
        self.port = int(getattr(self._server, 'server_port', port) or port)
    
    def run(self):
        self._server.serve_forever()
    
    def shutdown(self):
        try:
            self._server.shutdown()
        except Exception:
            pass


def wait_for_server(host, port, timeout=5.0):
    """Wait for server to be ready."""
    deadline = time.time() + timeout
    while time.time() < deadline:
        try:
            with socket.create_connection((host, port), timeout=0.3):
                return True
        except OSError:
            time.sleep(0.05)
    return False


# --- PyQt6 Bridge (exposed to JavaScript) ---
class JsBridge(QObject):
    """Bridge object exposed to JavaScript via QWebChannel."""
    
    # Signals for async operations
    fileDialogResult = pyqtSignal(str)
    viewerWindowRequested = pyqtSignal(str, str)  # filepath, title
    closeConfirmed = pyqtSignal(bool)  # For unsaved changes check
    keyboardToggleRequested = pyqtSignal()      # Toggle on-screen keyboard
    openDialogRequested = pyqtSignal(str)       # Open a named popup dialog
    closeDialogRequested = pyqtSignal(str)      # Close (hide) a named popup dialog
    showLoaderRequested = pyqtSignal(str, str)  # Show blocking loader (title, message)
    hideLoaderRequested = pyqtSignal()          # Hide blocking loader
    
    def __init__(self, main_window):
        super().__init__()
        self.main_window = main_window
    
    @pyqtSlot(result=str)
    def openFileDialog(self):
        """Open native file dialog."""
        filepath, _ = QFileDialog.getOpenFileName(
            self.main_window,
            "Open Document",
            LIBRARY_DIR,
            "śikṣāmitra Documents (*.smdoc);;Word Documents (*.docx);;HTML Files (*.html *.htm);;All Files (*)"
        )
        return filepath or ''

    @pyqtSlot(str, result=str)
    def saveFileDialog(self, default_name):
        """Open native save dialog."""
        # Strip any existing extension from the filename
        # The file type filter will determine the final extension
        base_name = default_name or 'document'
        
        # Determine original format to set the appropriate filter order
        is_html = base_name.lower().endswith(('.html', '.htm'))
        is_smdoc = base_name.lower().endswith('.smdoc')
        
        # Strip extension for the filename field
        if is_html:
            base_name = re.sub(r'\.html?$', '', base_name, flags=re.IGNORECASE)
            # HTML files default to HTML filter first
            filter_str = "HTML Files (*.html *.htm);;śikṣāmitra Documents (*.smdoc);;All Files (*)"
        elif is_smdoc:
            base_name = re.sub(r'\.smdoc$', '', base_name, flags=re.IGNORECASE)
            filter_str = "śikṣāmitra Documents (*.smdoc);;HTML Files (*.html *.htm);;All Files (*)"
        else:
            # For new/unknown files, default to .smdoc
            filter_str = "śikṣāmitra Documents (*.smdoc);;HTML Files (*.html *.htm);;All Files (*)"
        
        default_path = os.path.join(LIBRARY_DIR, base_name)
        
        filepath, selected_filter = QFileDialog.getSaveFileName(
            self.main_window,
            "Save Document",
            default_path,
            filter_str
        )
        return filepath or ''
    
    @pyqtSlot(str)
    def setWindowTitle(self, title):
        """Set main window title."""
        self.main_window.setWindowTitle(title)
    
    @pyqtSlot(str, str)
    def openViewerWindow(self, filepath, title):
        """Open document in viewer window."""
        self.viewerWindowRequested.emit(filepath, title)
    
    @pyqtSlot(result=str)
    def importFileDialog(self):
        """Open file dialog for import."""
        filepath, _ = QFileDialog.getOpenFileName(
            self.main_window,
            "Import Document to Library",
            "",
            "śikṣāmitra Documents (*.smdoc);;Word Documents (*.docx);;HTML Files (*.html *.htm);;All Files (*)"
        )
        return filepath or ''
    
    @pyqtSlot(str)
    def openInExplorer(self, filepath):
        """Open file location in explorer."""
        import subprocess
        if not filepath:
            return

        try:
            if os.path.isdir(filepath):
                subprocess.Popen(['explorer', filepath])
            elif os.path.exists(filepath):
                subprocess.Popen(['explorer', '/select,', filepath])
        except Exception:
            pass
    
    @pyqtSlot(bool)
    def confirmClose(self, can_close):
        """Called from JS to confirm if window can close."""
        self.closeConfirmed.emit(can_close)

    @pyqtSlot()
    def toggleKeyboard(self):
        """Toggle the floating on-screen keyboard window."""
        self.keyboardToggleRequested.emit()

    @pyqtSlot(str)
    def openDialog(self, dialog_id: str):
        """Open (or raise) a named popup dialog window."""
        self.openDialogRequested.emit(dialog_id)

    @pyqtSlot(str)
    def closeDialog(self, dialog_id: str):
        """Hide a named popup dialog window (window.close() is blocked in QtWebEngine)."""
        self.closeDialogRequested.emit(dialog_id)

    @pyqtSlot(str, str)
    def showLoader(self, title: str, message: str):
        """Show the blocking loader OS window with the given title and message."""
        self.showLoaderRequested.emit(title, message)

    @pyqtSlot()
    def hideLoader(self):
        """Hide the blocking loader OS window."""
        self.hideLoaderRequested.emit()


# --- Document Viewer Window ---
class ViewerWindow(QMainWindow):
    """Standalone viewer window for documents."""
    
    def __init__(self, filepath, title, server_url, theme='light', parent=None):
        super().__init__(parent)
        self.filepath = filepath
        self._theme = theme
        
        # Window setup
        self.setWindowTitle(f"{title} — śikṣāmitra")
        self.resize(900, 700)
        
        # Set icon
        if os.path.exists(ICON_PATH):
            self.setWindowIcon(QIcon(ICON_PATH))
        elif os.path.exists(LOGO_PATH):
            self.setWindowIcon(QIcon(LOGO_PATH))
        
        # Central widget
        central = QWidget()
        self.setCentralWidget(central)
        layout = QVBoxLayout(central)
        layout.setContentsMargins(0, 0, 0, 0)

        # WebView
        self.web_view = QWebEngineView()
        self.web_view.setContextMenuPolicy(Qt.ContextMenuPolicy.NoContextMenu)
        
        # Enable PDF viewer and plugins
        settings = self.web_view.settings()
        settings.setAttribute(QWebEngineSettings.WebAttribute.PluginsEnabled, True)
        settings.setAttribute(QWebEngineSettings.WebAttribute.PdfViewerEnabled, True)
        settings.setAttribute(QWebEngineSettings.WebAttribute.LocalContentCanAccessRemoteUrls, True)
        
        layout.addWidget(self.web_view)

        # Loading overlay (native, does not modify the HTML being viewed)
        self._overlay = QWidget(self.web_view)
        self._overlay.setAttribute(Qt.WidgetAttribute.WA_StyledBackground, True)
        self._overlay.setStyleSheet(
            "QWidget { background: rgba(0,0,0,0.25); }"
        )
        overlay_layout = QVBoxLayout(self._overlay)
        overlay_layout.setContentsMargins(0, 0, 0, 0)
        overlay_layout.setSpacing(0)

        overlay_card = QWidget(self._overlay)
        overlay_card.setStyleSheet(
            "QWidget { background: rgba(255,255,255,0.92); border-radius: 14px; }"
        )
        card_layout = QVBoxLayout(overlay_card)
        card_layout.setContentsMargins(22, 18, 22, 18)
        card_layout.setSpacing(8)

        title_lbl = QLabel("Opening…")
        title_lbl.setStyleSheet("QLabel { font-size: 14px; font-weight: 700; color: #222; }")
        doc_name = os.path.basename(filepath) if filepath else ""
        # For .smdoc files, show the actual title without extension
        if doc_name.lower().endswith('.smdoc'):
            doc_name = doc_name[:-6]  # Remove .smdoc extension for cleaner display
        msg_lbl = QLabel(doc_name)
        msg_lbl.setStyleSheet("QLabel { font-size: 11px; color: #555; }")
        msg_lbl.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)

        card_layout.addWidget(title_lbl)
        card_layout.addWidget(msg_lbl)

        overlay_layout.addStretch(1)
        overlay_layout.addWidget(overlay_card, alignment=Qt.AlignmentFlag.AlignHCenter)
        overlay_layout.addStretch(1)

        self._overlay.show()

        # Hook load signals
        self.web_view.loadStarted.connect(self._show_overlay)
        self.web_view.loadFinished.connect(self._hide_overlay)

        # Load raw HTML (no injected toolbar/quill wrappers)
        self._load_document(filepath)
        
        # Show maximized (fullscreen)
        self.showMaximized()

    def resizeEvent(self, event):
        super().resizeEvent(event)
        try:
            if getattr(self, '_overlay', None):
                self._overlay.setGeometry(0, 0, self.web_view.width(), self.web_view.height())
        except Exception:
            pass

    def _show_overlay(self):
        try:
            if getattr(self, '_overlay', None):
                self._overlay.setGeometry(0, 0, self.web_view.width(), self.web_view.height())
                self._overlay.show()
        except Exception:
            pass

    def _hide_overlay(self, ok=True):
        try:
            if getattr(self, '_overlay', None):
                self._overlay.hide()
        except Exception:
            pass
        
        # Inject theme CSS after document loads
        if ok and hasattr(self, '_theme') and self._theme == 'dark':
            self._inject_dark_theme()
    
    def _inject_dark_theme(self):
        """Inject dark theme styles into the loaded document."""
        dark_css = """
        (function() {
            var style = document.createElement('style');
            style.id = 'viewer-dark-theme';
            style.textContent = `
                html, body {
                    background: #1a1a1a !important;
                    color: #e8e8e8 !important;
                }
                body * {
                    color: inherit;
                }
            `;
            document.head.appendChild(style);
            document.body.setAttribute('data-theme', 'dark');
        })();
        """
        self.web_view.page().runJavaScript(dark_css)
    
    def _load_document(self, filepath):
        """Load document content with appropriate viewer."""
        try:
            if not filepath or not os.path.exists(filepath):
                self._show_error("File not found", filepath)
                return

            ext = os.path.splitext(filepath)[1].lower()
            
            # .smdoc format - needs to be hydrated to full HTML
            if ext == '.smdoc':
                self._load_smdoc_document(filepath)
            # Image formats
            elif ext in ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.svg', '.webp', '.ico'):
                self._load_image_document(filepath)
            # Text formats
            elif ext in ('.txt', '.md', '.py', '.js', '.css', '.json', '.xml', '.log', '.ini', '.conf', '.sh', '.bat', '.csv'):
                self._load_text_document(filepath)
            # Default (HTML, PDF, etc.)
            else:
                # Load as local file URL so rendering matches a normal browser view.
                self.web_view.setUrl(QUrl.fromLocalFile(os.path.abspath(filepath)))
        except Exception as e:
            self._show_error(str(e), filepath)
    
    def _load_smdoc_document(self, filepath):
        """Load .smdoc file and hydrate to full HTML for viewing."""
        try:
            print(f"[Viewer] Loading .smdoc file: {filepath}")
            
            # Read the file - may be compressed
            with open(filepath, 'rb') as f:
                raw_data = f.read()
            
            print(f"[Viewer] Read {len(raw_data):,} bytes from file")
            
            # Decompress if needed (handles both compressed and uncompressed)
            content_str = decompress_smdoc(raw_data)
            
            print(f"[Viewer] Decompressed to {len(content_str):,} bytes")
            
            smdoc = json.loads(content_str)
            
            # Handle double-encoded JSON (backward compatibility)
            if isinstance(smdoc, str):
                print("[Viewer] SMDoc was double-encoded, parsing inner JSON")
                smdoc = json.loads(smdoc)
            
            content = smdoc.get('content', '')
            print(f"[Viewer] Content length: {len(content):,} bytes")
            meta = smdoc.get('meta', {})
            styles = smdoc.get('styles', {})
            title = meta.get('title', os.path.splitext(os.path.basename(filepath))[0])
            theme = styles.get('theme', 'light')
            
            # Load embedded font data
            font_data_uri = ''
            try:
                font_path = os.path.join(os.path.dirname(__file__), 'palladio_font_data_uri.txt')
                if os.path.exists(font_path):
                    with open(font_path, 'r', encoding='utf-8') as f:
                        font_content = f.read()
                    # The font file uses JS string concatenation like 'data...' + 'more...' + '...'
                    # First, find all string literals and concatenate them
                    # Find all quoted strings (both single and double quotes)
                    string_parts = re.findall(r"'([^']*)'", font_content)
                    if string_parts:
                        # Join all string parts
                        full_string = ''.join(string_parts)
                        # Check if it looks like a data URI
                        if full_string.startswith('data:font/'):
                            font_data_uri = full_string
            except Exception as e:
                print(f"Failed to load font data: {e}")
            
            # Generate full HTML
            print(f"[Viewer] Generating HTML with theme: {theme}")
            html = self._generate_smdoc_html(content, title, theme, font_data_uri)
            print(f"[Viewer] Generated HTML: {len(html):,} bytes")
            
            # For large files, Qt's setHtml() can fail silently.
            # Write to a temp file and load via URL for better reliability.
            import tempfile
            temp_dir = os.path.join(CACHE_DIR, 'viewer_temp')
            os.makedirs(temp_dir, exist_ok=True)
            
            # Use a unique temp filename
            temp_filename = f"_viewer_{os.getpid()}.html"
            temp_path = os.path.join(temp_dir, temp_filename)
            
            print(f"[Viewer] Writing to temp file: {temp_path}")
            with open(temp_path, 'w', encoding='utf-8') as f:
                f.write(html)
            
            # Load from file URL - more reliable for large documents
            print(f"[Viewer] Loading from file URL...")
            self.web_view.setUrl(QUrl.fromLocalFile(temp_path))
            print(f"[Viewer] setUrl() called successfully")
            
        except Exception as e:
            print(f"[Viewer] ERROR: {str(e)}")
            import traceback
            traceback.print_exc()
            self._show_error(f"Failed to load .smdoc file: {str(e)}", filepath)
    
    def _generate_smdoc_html(self, content, title, theme, font_data_uri=''):
        """Generate full HTML from .smdoc content - matching the full exported HTML format."""
        # This matches the getDocumentHTML() function from editor-quill.js
        # It includes the full side menu, theme toggle, audio controls, and all styles
        
        embedded_styles = f'''
        /* Embedded URW Palladio ITU font */
        @font-face {{
            font-family: 'URW Palladio ITU';
            font-style: normal;
            font-weight: 400;
            font-display: swap;
            src: url('{font_data_uri}') format('truetype');
        }}

        :root {{
            --primary: #b8813d;
            --primary-hover: #9a6b2f;
            --accent: #d97706;
            --bg-body: #f2f2f5;
            --bg-surface: #ffffff;
            --text-primary: #1b1b1f;
            --text-secondary: #5c5c66;
            --border: #dddde1;
            --border-subtle: #e8e8ec;
            --holding-border: #10b981;
            --holding-long: #ef4444;
            --translation-text-color: #5c5c66;
            --shadow: 0 4px 12px rgba(0,0,0,0.07);
            --shadow-lg: 0 12px 32px rgba(0,0,0,0.10);
        }}

        [data-theme='dark'] {{
            --primary: #d4a574;
            --primary-hover: #e6b886;
            --accent: #ff9933;
            --bg-body: #111113;
            --bg-surface: #1c1c1f;
            --text-primary: #e8e8ec;
            --text-secondary: #a0a0a8;
            --border: #333338;
            --border-subtle: #2a2a2e;
            --holding-border: #34d399;
            --holding-long: #f87171;
            --translation-text-color: #a0a0a8;
            --shadow: 0 4px 12px rgba(0,0,0,0.4);
            --shadow-lg: 0 12px 32px rgba(0,0,0,0.5);
        }}

        * {{ margin: 0; padding: 0; box-sizing: border-box; }}

        body {{
            margin: 0;
            padding: 3rem 1.5rem;
            background: var(--bg-body);
            color: var(--text-primary);
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Roboto', 'Helvetica Neue', Arial, sans-serif;
            font-size: 28px;
            line-height: 1.55;
            min-height: 100vh;
            transition: background-color 0.3s ease, color 0.3s ease;
        }}

        .paper {{
            max-width: 960px;
            margin: 0 auto;
            background: var(--bg-surface);
            border-radius: 18px;
            box-shadow: var(--shadow-lg);
            border: 1px solid var(--border);
            padding: 3rem;
            transition: background-color 0.3s ease, color 0.3s ease, border-color 0.3s ease;
        }}

        .content {{
            min-height: 60vh;
        }}

        .ql-editor {{
            padding: 0;
            color: var(--text-primary);
            background: transparent;
            line-height: 1.5;
            font-size: 28px;
            font-family: Arial, sans-serif;
            overflow-wrap: break-word;
            word-break: normal;
            white-space: pre-wrap;
        }}

        .ql-editor strong, strong, .ql-bold {{ font-weight: 700; }}
        .ql-editor em, em, .ql-italic {{ font-style: italic; }}
        .ql-editor u, u, .ql-underline {{ text-decoration: underline; }}
        .ql-editor s, s, .ql-strike {{ text-decoration: line-through; }}

        .ql-font-gentium,
        [style*="font-family: 'Gentium Plus'"],
        [style*='font-family: "Gentium Plus"'],
        [style*="font-family: Gentium Plus"],
        [style*="font-family: 'URW Palladio ITU'"],
        [style*='font-family: "URW Palladio ITU"'],
        [style*="font-family: URW Palladio ITU"] {{
            font-family: 'URW Palladio ITU', serif !important;
        }}

        .ql-font-arial, [style*="font-family: Arial"] {{ font-family: Arial, sans-serif !important; }}
        .ql-font-times {{ font-family: 'Times New Roman', serif !important; }}
        .ql-font-calibri {{ font-family: Calibri, sans-serif !important; }}
        .ql-font-georgia {{ font-family: Georgia, serif !important; }}
        .ql-font-verdana {{ font-family: Verdana, sans-serif !important; }}
        .ql-font-courier {{ font-family: 'Courier New', monospace !important; }}
        .ql-font-noto {{ font-family: 'Noto Sans Devanagari', sans-serif !important; }}

        .ql-size-8px {{ font-size: 8px; }}
        .ql-size-9px {{ font-size: 9px; }}
        .ql-size-10px {{ font-size: 10px; }}
        .ql-size-11px {{ font-size: 11px; }}
        .ql-size-12px {{ font-size: 12px; }}
        .ql-size-14px {{ font-size: 14px; }}
        .ql-size-16px {{ font-size: 16px; }}
        .ql-size-18px {{ font-size: 18px; }}
        .ql-size-20px {{ font-size: 20px; }}
        .ql-size-22px {{ font-size: 22px; }}
        .ql-size-24px {{ font-size: 24px; }}
        .ql-size-28px {{ font-size: 28px; }}
        .ql-size-32px {{ font-size: 32px; }}
        .ql-size-36px {{ font-size: 36px; }}
        .ql-size-48px {{ font-size: 48px; }}
        .ql-size-72px {{ font-size: 72px; }}

        .ql-align-left {{ text-align: left; }}
        .ql-align-center {{ text-align: center; }}
        .ql-align-right {{ text-align: right; }}
        .ql-align-justify {{ text-align: justify; }}

        .ql-holding-short {{
            border: 1px solid var(--holding-border);
            border-radius: 3px;
            padding: 0;
            display: inline;
            box-decoration-break: clone;
            -webkit-box-decoration-break: clone;
        }}

        .ql-holding-long {{
            border: 2px solid var(--holding-border);
            border-radius: 3px;
            padding: 0;
            display: inline;
            box-decoration-break: clone;
            -webkit-box-decoration-break: clone;
        }}

        .ql-change-style {{
            font-style: italic;
            color: #1d4ed8;
            white-space: pre;
        }}

        [data-theme='dark'] .ql-change-style {{ color: #60a5fa; }}

        .ql-translation-style {{
            font-style: italic;
            color: var(--translation-text-color);
            white-space: pre-wrap;
            font-size: 0.8em;
            line-height: 1.35;
            display: inline;
        }}

        .ql-doc-translation {{
            font-style: italic;
            color: var(--translation-text-color);
            white-space: pre-wrap;
            font-size: 0.8em;
            line-height: 1.35;
            display: block;
            margin: 0.55em 0;
        }}

        .ql-editor br.soft-break, br.soft-break {{
            display: inline;
            margin: 0;
            line-height: inherit;
        }}

        .content p {{ position: relative; }}

        .ql-audio-attachment {{
            display: inline !important;
            width: 0 !important;
            height: 0 !important;
            overflow: hidden !important;
            margin: 0 !important;
            padding: 0 !important;
            font-size: 0 !important;
            line-height: 0 !important;
        }}

        .ql-audio-attachment audio {{ display: none !important; }}

        .audio-play-button {{
            position: absolute;
            left: -60px;
            top: 0;
            width: 40px;
            height: 40px;
            min-width: 40px;
            border-radius: 50%;
            border: 2px solid var(--primary);
            background: var(--bg-surface);
            color: var(--primary);
            cursor: pointer;
            display: flex !important;
            align-items: center;
            justify-content: center;
            transition: all 0.2s ease;
            box-shadow: var(--shadow);
            padding: 0;
            margin: 0;
            z-index: 10;
        }}

        .audio-play-button:hover {{
            background: var(--primary);
            color: #ffffff;
            transform: scale(1.1);
        }}

        .audio-icon {{
            display: flex;
            align-items: center;
            justify-content: center;
            width: 18px;
            height: 18px;
        }}

        .audio-icon svg {{
            width: 16px;
            height: 16px;
            stroke: currentColor;
            fill: none;
            stroke-width: 2;
            stroke-linecap: round;
            stroke-linejoin: round;
        }}

        .audio-play-button[data-state='play'] .audio-icon .icon-play {{ display: block; }}
        .audio-play-button[data-state='play'] .audio-icon .icon-stop {{ display: none; }}
        .audio-play-button[data-state='pause'] .audio-icon .icon-play {{ display: none; }}
        .audio-play-button[data-state='pause'] .audio-icon .icon-stop {{ display: block; }}

        .ql-comment-style {{
            font-style: italic;
            color: #92400e;
            background: rgba(251, 191, 36, 0.25);
            border-left: 3px solid #f59e0b;
            padding: 0.1em 0.35em;
            border-radius: 4px;
            display: inline;
            white-space: pre-wrap;
            box-decoration-break: clone;
            -webkit-box-decoration-break: clone;
        }}

        [data-theme='dark'] .ql-comment-style {{
            color: #fcd34d;
            background: rgba(146, 64, 14, 0.35);
            border-left-color: #fbbf24;
        }}

        body.hide-comments .ql-comment-style {{ display: none !important; }}

        .ql-short-pause {{ color: #2563eb; font-weight: bold; white-space: pre; }}
        .ql-long-pause {{ color: #dc2626; font-weight: bold; white-space: pre; }}

        .ql-dirgha-char, .ql-dirgha, .ql-dirgha-true {{
            --dirgha-line-color: #1d4ed8;
            --dirgha-line-thickness: 0.08em;
            --dirgha-line-y: 0.28em;
            --dirgha-line-extend: 0.12em;
            color: inherit !important;
            text-decoration: none !important;
            position: relative !important;
            display: inline-block !important;
        }}

        .ql-dirgha-char::before, .ql-dirgha::before, .ql-dirgha-true::before {{
            content: '';
            position: absolute;
            left: calc(-1 * var(--dirgha-line-extend));
            right: calc(-1 * var(--dirgha-line-extend));
            top: var(--dirgha-line-y);
            height: var(--dirgha-line-thickness);
            background: var(--dirgha-line-color);
            pointer-events: none;
        }}

        [data-theme='dark'] .ql-dirgha-char, [data-theme='dark'] .ql-dirgha, [data-theme='dark'] .ql-dirgha-true {{
            --dirgha-line-color: #60a5fa;
        }}

        .ql-svara-char, .ql-svara, .ql-svara-true {{
            font-family: "URW Palladio ITU", "Times New Roman", serif !important;
            font-size: inherit !important;
            font-weight: 800 !important;
            color: #cc1b1b !important;
            font-style: normal !important;
            text-decoration: none !important;
            display: inline !important;
        }}

        [data-theme='dark'] .ql-svara-char, [data-theme='dark'] .ql-svara, [data-theme='dark'] .ql-svara-true {{
            color: #ff4d4d !important;
        }}

        p {{ margin: 0; padding: 0; }}
        p + p {{ margin-top: 0.85em; }}

        /* Document paragraph styles - Compact Ochre Design */
        .ql-doc-title {{
            font-size: 1.75em;
            font-weight: 600;
            letter-spacing: -0.01em;
            color: #92400e;
            text-align: left;
            margin: 0.5em 0 0.3em 0;
            line-height: 1.25;
            border-bottom: 2px solid #d97706;
            padding-bottom: 0.2em;
        }}

        [data-theme='dark'] .ql-doc-title {{
            color: #fbbf24;
            border-bottom-color: #d97706;
        }}

        .ql-doc-subtitle {{
            font-size: 0.85em;
            font-weight: 500;
            color: #78716c;
            text-align: left;
            margin: 0 0 0.8em 0;
            line-height: 1.3;
            letter-spacing: 0.04em;
            text-transform: uppercase;
        }}

        [data-theme='dark'] .ql-doc-subtitle {{
            color: #a8a29e;
        }}

        .ql-doc-section {{
            font-size: 1.15em;
            font-weight: 600;
            color: #b45309;
            margin: 1em 0 0.3em 0;
            line-height: 1.3;
            border-left: 3px solid #d97706;
            padding-left: 0.5em;
        }}

        [data-theme='dark'] .ql-doc-section {{
            color: #fbbf24;
            border-left-color: #f59e0b;
        }}

        .ql-doc-subsection {{
            font-size: 0.9em;
            font-weight: 600;
            color: #78716c;
            margin: 0.8em 0 0.2em 0;
            line-height: 1.3;
            letter-spacing: 0.02em;
            border-left: 2px solid #d4a574;
            padding-left: 0.4em;
        }}

        [data-theme='dark'] .ql-doc-subsection {{
            color: #a8a29e;
            border-left-color: #d97706;
        }}

        /* Theme toggle button */
        .theme-toggle {{
            position: fixed;
            top: 20px;
            right: 20px;
            width: 40px;
            height: 40px;
            border: 1px solid var(--border);
            border-radius: 50%;
            background: var(--bg-surface);
            color: var(--text-primary);
            cursor: pointer;
            font-size: 20px;
            display: flex;
            align-items: center;
            justify-content: center;
            transition: all 0.2s;
            box-shadow: var(--shadow);
            z-index: 1000;
        }}

        .theme-toggle:hover {{
            transform: scale(1.1);
            box-shadow: var(--shadow-lg);
        }}

        /* Options Panel Styles */
        .options-toggle-btn {{
            position: fixed;
            top: 20px;
            left: 20px;
            width: 44px;
            height: 44px;
            border: 1px solid var(--border);
            border-radius: 50%;
            background: var(--bg-surface);
            color: var(--text-primary);
            cursor: pointer;
            font-size: 20px;
            display: flex;
            align-items: center;
            justify-content: center;
            transition: all 0.2s;
            box-shadow: var(--shadow);
            z-index: 2000;
        }}

        .options-toggle-btn:hover {{
            transform: scale(1.1);
            box-shadow: var(--shadow-lg);
            background: var(--primary);
            border-color: var(--primary);
            color: #ffffff;
        }}

        .options-toggle-btn.hidden {{ opacity: 0; pointer-events: none; }}
        .options-icon {{ display: block; line-height: 1; }}

        .options-container {{
            position: fixed;
            top: 0;
            left: 0;
            width: 360px;
            height: 100vh;
            background: var(--bg-surface);
            border-right: 1px solid var(--border);
            box-shadow: var(--shadow-lg);
            z-index: 1900;
            transform: translateX(-100%);
            transition: transform 0.3s ease;
            display: flex;
            flex-direction: column;
            overflow: hidden;
        }}

        .options-container.open {{ transform: translateX(0); }}

        .options-header {{
            padding: 20px;
            border-bottom: 1px solid var(--border);
            background: var(--bg-surface);
            display: flex;
            justify-content: space-between;
            align-items: center;
            flex-shrink: 0;
        }}

        .options-title {{
            font-size: 11px;
            font-weight: 600;
            color: var(--text-secondary);
            text-transform: uppercase;
            letter-spacing: 0.5px;
            margin: 0;
        }}

        .options-close-btn {{
            width: 28px;
            height: 28px;
            border: none;
            border-radius: 6px;
            background: transparent;
            color: var(--text-secondary);
            cursor: pointer;
            font-size: 18px;
            display: flex;
            align-items: center;
            justify-content: center;
            transition: all 0.15s;
        }}

        .options-close-btn:hover {{ background: var(--bg-body); color: var(--text-primary); }}

        .options-content {{
            flex: 1;
            overflow-y: auto;
            padding: 16px 20px;
        }}

        .option-group {{ margin-bottom: 20px; }}
        .option-group:last-child {{ margin-bottom: 0; }}

        .option-group-title {{
            font-size: 11px;
            font-weight: 600;
            color: var(--text-secondary);
            margin: 0 0 10px 0;
            text-transform: uppercase;
            letter-spacing: 0.5px;
        }}

        /* option-btn kept for any non-switch uses (e.g. options toggle) */

        /* Switch toggle — matches editor's switch-checkbox */
        .option-switch {{
            display: flex;
            align-items: center;
            gap: 12px;
            padding: 8px 0;
            cursor: pointer;
            user-select: none;
        }}

        .option-switch input {{ display: none; }}

        .option-switch .switch-track {{
            width: 40px;
            height: 22px;
            background: var(--border);
            border-radius: 22px;
            position: relative;
            transition: background 0.2s;
            flex-shrink: 0;
        }}

        .option-switch .switch-knob {{
            position: absolute;
            width: 16px;
            height: 16px;
            left: 3px;
            top: 3px;
            background: white;
            border-radius: 50%;
            transition: transform 0.2s;
        }}

        .option-switch.on .switch-track {{
            background: var(--primary);
        }}

        .option-switch.on .switch-knob {{
            transform: translateX(18px);
        }}

        .option-switch .switch-label {{
            font-size: 14px;
            font-weight: 500;
            color: var(--text-primary);
            display: flex;
            align-items: center;
            gap: 8px;
        }}

        .option-switch .theme-icon {{
            color: var(--text-secondary);
            flex-shrink: 0;
        }}

        .option-group-bottom {{
            margin-top: auto;
            padding-top: 16px;
            border-top: 1px solid var(--border-subtle);
        }}

        /* Speed Controls */
        .speed-controls {{
            display: flex;
            gap: 12px;
            align-items: center;
            margin-bottom: 8px;
        }}

        .speed-label {{
            font-size: 13px;
            font-weight: 600;
            color: var(--text-secondary);
            min-width: 32px;
        }}

        .speed-slider {{
            flex: 1;
            height: 6px;
            border-radius: 3px;
            background: var(--border);
            outline: none;
            -webkit-appearance: none;
        }}

        .speed-slider::-webkit-slider-thumb {{
            -webkit-appearance: none;
            width: 20px;
            height: 20px;
            border-radius: 50%;
            background: var(--primary);
            cursor: pointer;
        }}

        .speed-display {{
            text-align: center;
            font-size: 14px;
            font-weight: 600;
            color: var(--primary);
            padding: 4px 0;
        }}

        /* ToC styles */
        .toc-nav {{ max-height: none; padding: 0; }}
        .toc-list {{ list-style: none; padding: 0; margin: 0; }}
        .toc-item {{ margin: 0; padding: 0; }}

        .toc-link {{
            display: block;
            padding: 6px 10px;
            color: var(--text-primary);
            text-decoration: none;
            font-size: 13px;
            line-height: 1.4;
            border-radius: 6px;
            transition: background 0.1s;
            cursor: pointer;
        }}

        .toc-link:hover {{ background: var(--bg-body); }}
        .toc-link.active {{
            background: rgba(184, 129, 61, 0.08);
            color: var(--primary);
        }}

        .toc-link.level-title {{ font-size: 13px; font-weight: 600; }}
        .toc-link.level-subtitle {{ font-size: 13px; font-weight: 500; padding-left: 16px; }}
        .toc-link.level-section {{ font-size: 12px; padding-left: 28px; }}
        .toc-link.level-subsection {{ font-size: 12px; color: var(--text-secondary); padding-left: 40px; }}

        .options-overlay {{ display: none; }}

        body.hide-translations .ql-translation-style,
        body.hide-translations .ql-doc-translation {{ display: none !important; }}
        body.hide-translations .translation-break {{ display: none !important; }}
        body.hide-audio .audio-play-button {{ display: none !important; }}

        @media (max-width: 767px) {{
            .options-toggle-btn {{ width: 48px; height: 48px; font-size: 22px; }}
            .options-container {{ width: 100%; max-width: 100vw; border-right: none; }}
            .options-overlay {{
                display: none;
                position: fixed;
                top: 0; left: 0; right: 0; bottom: 0;
                background: rgba(0, 0, 0, 0.5);
                z-index: 1800;
                opacity: 0;
                transition: opacity 0.3s ease;
            }}
            .options-overlay.visible {{ display: block; opacity: 1; }}
        }}

        @media print {{
            body {{ background: #ffffff !important; padding: 0; }}
            .paper {{ box-shadow: none; border-radius: 0; border: none; margin: 0; max-width: 100%; padding: 2.5cm; }}
            .theme-toggle, .options-toggle-btn, .options-container, .options-overlay {{ display: none !important; }}
        }}
        '''
        
        return f'''<!DOCTYPE html>
<html lang="sa">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title} — śikṣāmitra</title>
    <style>
{embedded_styles}
    </style>
</head>
<body data-theme="{theme}">
    <!-- Options Toggle Button -->
    <button class="options-toggle-btn" id="optionsToggleBtn" onclick="toggleOptions()" title="Document Options" aria-label="Document Options">
        <svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><line x1="3" y1="6" x2="21" y2="6"/><line x1="3" y1="12" x2="21" y2="12"/><line x1="3" y1="18" x2="21" y2="18"/></svg>
    </button>

    <!-- Options Panel -->
    <div class="options-container" id="optionsContainer">
        <div class="options-header">
            <h2 class="options-title">Options</h2>
            <button class="options-close-btn" onclick="toggleOptions()" aria-label="Close Options">
                <svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><line x1="18" y1="6" x2="6" y2="18"/><line x1="6" y1="6" x2="18" y2="18"/></svg>
            </button>
        </div>

        <div class="options-content">
            <div class="option-group">
                <h3 class="option-group-title">Visibility</h3>
                <label class="option-switch on" onclick="toggleTranslations()" id="translationsSwitch">
                    <span class="switch-track"><span class="switch-knob"></span></span>
                    <span class="switch-label">Translations</span>
                </label>
                <label class="option-switch on" onclick="toggleAudio()" id="audioSwitch">
                    <span class="switch-track"><span class="switch-knob"></span></span>
                    <span class="switch-label">Audio Buttons</span>
                </label>
                <label class="option-switch on" onclick="toggleComments()" id="commentsSwitch">
                    <span class="switch-track"><span class="switch-knob"></span></span>
                    <span class="switch-label">Comments</span>
                </label>
            </div>

            <div class="option-group">
                <h3 class="option-group-title">Audio Playback Speed</h3>
                <div class="speed-controls">
                    <span class="speed-label">0.5x</span>
                    <input type="range" class="speed-slider" id="speedSlider" min="0.5" max="2.0" step="0.1" value="1.0" oninput="setPlaybackSpeed(parseFloat(this.value))">
                    <span class="speed-label">2x</span>
                </div>
                <div class="speed-display" id="speedDisplay">Speed: 1.0x</div>
            </div>
            
            <div class="option-group">
                <h3 class="option-group-title">Navigation</h3>
                <nav class="toc-nav" id="tocNav"></nav>
            </div>

            <div class="option-group option-group-bottom">
                <label class="option-switch" onclick="toggleTheme()" id="themeSwitch">
                    <span class="switch-track"><span class="switch-knob"></span></span>
                    <span class="switch-label">
                        <svg class="theme-icon" id="themeIconSvg" width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M21 12.79A9 9 0 1 1 11.21 3 7 7 0 0 0 21 12.79z"/></svg>
                        <span id="themeText">Dark Mode</span>
                    </span>
                </label>
            </div>
        </div>
    </div>
    
    <div class="options-overlay" id="optionsOverlay" onclick="toggleOptions()"></div>
    
    <div class="paper">
        <div class="content ql-editor">{content}</div>
    </div>
    
    <script>
        function initAudioButtons() {{
            const attachments = Array.from(document.querySelectorAll('.ql-audio-attachment'));
            if (!attachments.length) return;
            let activeAudio = null;
            
            attachments.forEach((attachment) => {{
                const audio = attachment.querySelector('audio');
                if (!audio || !audio.src) return;
                const paragraph = attachment.closest('p');
                if (!paragraph) return;
                
                const startTime = parseFloat(attachment.dataset.startTime) || 0;
                const endTime = attachment.dataset.endTime ? parseFloat(attachment.dataset.endTime) : null;
                const label = attachment.dataset.audioLabel || 'Audio';
                
                const playButton = document.createElement('button');
                playButton.type = 'button';
                playButton.className = 'audio-play-button';
                playButton.setAttribute('aria-label', 'Play audio: ' + label);
                playButton.title = label;
                playButton.dataset.state = 'play';
                // SVG play and stop icons
                playButton.innerHTML = '<span class="audio-icon">' +
                    '<svg class="icon-play" viewBox="0 0 24 24"><path d="M16.6582 9.28638C18.098 10.1862 18.8178 10.6361 19.0647 11.2122C19.2803 11.7152 19.2803 12.2847 19.0647 12.7878C18.8178 13.3638 18.098 13.8137 16.6582 14.7136L9.896 18.94C8.29805 19.9387 7.49907 20.4381 6.83973 20.385C6.26501 20.3388 5.73818 20.0469 5.3944 19.584C5 19.053 5 18.1108 5 16.2264V7.77357C5 5.88919 5 4.94701 5.3944 4.41598C5.73818 3.9531 6.26501 3.66111 6.83973 3.6149C7.49907 3.5619 8.29805 4.06126 9.896 5.05998L16.6582 9.28638Z"/></svg>' +
                    '<svg class="icon-stop" viewBox="0 0 24 24"><rect x="4" y="4" width="16" height="16" rx="2"/></svg>' +
                    '</span>';
                paragraph.appendChild(playButton);
                
                playButton.addEventListener('click', (e) => {{
                    e.preventDefault();
                    e.stopPropagation();
                    
                    if (activeAudio && activeAudio !== audio) {{
                        activeAudio.pause();
                        activeAudio.currentTime = 0;
                        const otherAttachment = activeAudio.closest('.ql-audio-attachment');
                        const otherButton = otherAttachment ? otherAttachment.closest('p').querySelector('.audio-play-button') : null;
                        if (otherAttachment) otherAttachment.classList.remove('playing');
                        if (otherButton) otherButton.dataset.state = 'play';
                    }}
                    
                    if (audio.paused) {{
                        audio.currentTime = startTime;
                        playButton.dataset.state = 'loading';
                        audio.play().then(() => {{
                            attachment.classList.add('playing');
                            playButton.dataset.state = 'pause';
                            activeAudio = audio;
                        }}).catch(() => {{
                            playButton.dataset.state = 'play';
                        }});
                    }} else {{
                        audio.pause();
                        audio.currentTime = startTime;
                        attachment.classList.remove('playing');
                        playButton.dataset.state = 'play';
                        activeAudio = null;
                    }}
                }});
                
                audio.addEventListener('ended', () => {{
                    attachment.classList.remove('playing');
                    playButton.dataset.state = 'play';
                    if (activeAudio === audio) activeAudio = null;
                }});
                
                audio.addEventListener('timeupdate', () => {{
                    if (endTime !== null && audio.currentTime >= endTime) {{
                        audio.pause();
                        audio.currentTime = startTime;
                        attachment.classList.remove('playing');
                        playButton.dataset.state = 'play';
                        if (activeAudio === audio) activeAudio = null;
                    }}
                }});
            }});
        }}
        
        initAudioButtons();
        
        function toggleOptions() {{
            const container = document.getElementById('optionsContainer');
            const overlay = document.getElementById('optionsOverlay');
            const toggleBtn = document.getElementById('optionsToggleBtn');
            const isOpen = container.classList.contains('open');
            
            if (isOpen) {{
                container.classList.remove('open');
                overlay.classList.remove('visible');
                toggleBtn.classList.remove('hidden');
                document.body.style.overflow = '';
            }} else {{
                container.classList.add('open');
                overlay.classList.add('visible');
                toggleBtn.classList.add('hidden');
                if (window.innerWidth < 768) document.body.style.overflow = 'hidden';
            }}
        }}
        
        function updateThemeIcon(theme) {{
            var sw = document.getElementById('themeSwitch');
            var text = document.getElementById('themeText');
            var icon = document.getElementById('themeIconSvg');
            if (sw) {{
                if (theme === 'dark') sw.classList.add('on');
                else sw.classList.remove('on');
            }}
            if (text) text.textContent = theme === 'dark' ? 'Light Mode' : 'Dark Mode';
            if (icon) {{
                if (theme === 'dark') {{
                    icon.innerHTML = '<circle cx="12" cy="12" r="4"/><path d="M12 2v2M12 20v2M4.93 4.93l1.41 1.41M17.66 17.66l1.41 1.41M2 12h2M20 12h2M6.34 17.66l-1.41 1.41M19.07 4.93l-1.41 1.41"/>';
                }} else {{
                    icon.innerHTML = '<path d="M21 12.79A9 9 0 1 1 11.21 3 7 7 0 0 0 21 12.79z"/>';
                }}
            }}
        }}

        function applyTheme(theme) {{
            const resolved = theme === 'dark' ? 'dark' : 'light';
            document.body.setAttribute('data-theme', resolved);
            localStorage.setItem('theme', resolved);
            updateThemeIcon(resolved);
        }}

        function toggleTheme() {{
            var sw = document.getElementById('themeSwitch');
            var isDark = sw.classList.contains('on');
            if (isDark) {{
                sw.classList.remove('on');
                applyTheme('light');
            }} else {{
                sw.classList.add('on');
                applyTheme('dark');
            }}
        }}

        function toggleTranslations() {{
            var sw = document.getElementById('translationsSwitch');
            sw.classList.toggle('on');
            if (sw.classList.contains('on')) document.body.classList.remove('hide-translations');
            else document.body.classList.add('hide-translations');
        }}

        function toggleAudio() {{
            var sw = document.getElementById('audioSwitch');
            sw.classList.toggle('on');
            if (sw.classList.contains('on')) document.body.classList.remove('hide-audio');
            else document.body.classList.add('hide-audio');
        }}

        function toggleComments() {{
            var sw = document.getElementById('commentsSwitch');
            if (!sw) return;
            sw.classList.toggle('on');
            if (sw.classList.contains('on')) document.body.classList.remove('hide-comments');
            else document.body.classList.add('hide-comments');
        }}
        
        function setPlaybackSpeed(speed) {{
            const audioElements = document.querySelectorAll('audio');
            audioElements.forEach(audio => {{ audio.playbackRate = speed; }});
            const slider = document.getElementById('speedSlider');
            const display = document.getElementById('speedDisplay');
            if (slider) slider.value = speed;
            if (display) display.textContent = 'Speed: ' + speed.toFixed(1) + 'x';
        }}

        function buildTableOfContents() {{
            const contentDiv = document.querySelector('.content.ql-editor');
            if (!contentDiv) return;
            
            const headings = contentDiv.querySelectorAll('.ql-doc-title, .ql-doc-subtitle, .ql-doc-section, .ql-doc-subsection');
            if (headings.length === 0) {{
                document.getElementById('tocNav').innerHTML = '<p style="padding: 12px; color: var(--text-secondary); text-align: center; font-style: italic; font-size: 14px;">No headings found</p>';
                return;
            }}
            
            const tocList = document.createElement('ul');
            tocList.className = 'toc-list';
            
            headings.forEach((heading, index) => {{
                const headingId = 'heading-' + index;
                heading.id = headingId;
                
                let level = 'title';
                if (heading.classList.contains('ql-doc-subtitle')) level = 'subtitle';
                else if (heading.classList.contains('ql-doc-section')) level = 'section';
                else if (heading.classList.contains('ql-doc-subsection')) level = 'subsection';
                
                const headingText = heading.textContent.trim() || 'Untitled';
                
                const tocItem = document.createElement('li');
                tocItem.className = 'toc-item';
                
                const tocLink = document.createElement('a');
                tocLink.className = 'toc-link level-' + level;
                tocLink.textContent = headingText;
                tocLink.href = '#' + headingId;
                tocLink.setAttribute('data-heading-id', headingId);
                
                tocLink.addEventListener('click', (e) => {{
                    e.preventDefault();
                    const targetElement = document.getElementById(headingId);
                    if (targetElement) {{
                        targetElement.scrollIntoView({{ behavior: 'smooth', block: 'start' }});
                        if (window.innerWidth < 768) toggleOptions();
                    }}
                }});
                
                tocItem.appendChild(tocLink);
                tocList.appendChild(tocItem);
            }});
            
            document.getElementById('tocNav').innerHTML = '';
            document.getElementById('tocNav').appendChild(tocList);
        }}

        buildTableOfContents();
        
        const savedTheme = localStorage.getItem('theme') || '{theme}';
        applyTheme(savedTheme);
        
        document.addEventListener('keydown', (e) => {{
            if (e.key === 'Escape') {{
                const container = document.getElementById('optionsContainer');
                if (container.classList.contains('open')) toggleOptions();
            }}
        }});
    </script>
</body>
</html>'''

    def _load_image_document(self, filepath):
        """Load image with zoom/rotate controls."""
        try:
            # Convert image to base64 to embed it (avoids some local file permission issues in templates)
            # or just use file:// URL. file:// is better for large images.
            file_url = QUrl.fromLocalFile(os.path.abspath(filepath)).toString()
            
            html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>Image Viewer</title>
<style>
    body {{
        margin: 0;
        padding: 0;
        background: #f0f0f0;
        height: 100vh;
        display: flex;
        flex-direction: column;
        overflow: hidden;
        font-family: system-ui, -apple-system, sans-serif;
    }}
    body[data-theme='dark'] {{
        background: #1a1a1a;
        color: #e0e0e0;
    }}
    .toolbar {{
        padding: 10px;
        background: var(--bg-surface);
        border-bottom: 1px solid var(--border);
        display: flex;
        gap: 10px;
        justify-content: center;
        z-index: 10;
    }}
    .btn {{
        padding: 6px 12px;
        border: 1px solid var(--border);
        background: transparent;
        border-radius: 6px;
        cursor: pointer;
        font-size: 14px;
        color: var(--text-primary);
        transition: all 0.2s;
    }}
    .btn:hover {{ background: var(--bg-body); border-color: var(--primary); }}
    
    .viewport {{
        flex: 1;
        overflow: auto;
        display: flex;
        align-items: center;
        justify-content: center;
        position: relative;
    }}
    #img-container {{
        transition: transform 0.2s ease-out;
        transform-origin: center center;
        box-shadow: 0 5px 20px rgba(0,0,0,0.1);
    }}
    img {{
        max-width: 90vw;
        max-height: 90vh;
        display: block;
    }}
</style>
</head>
<body data-theme="{self._theme}">
    <div class="toolbar">
        <button class="btn" onclick="zoomIn()">Zoom In (+)</button>
        <button class="btn" onclick="zoomOut()">Zoom Out (-)</button>
        <button class="btn" onclick="rotateLeft()">Rotate L (↺)</button>
        <button class="btn" onclick="rotateRight()">Rotate R (↻)</button>
        <button class="btn" onclick="reset()">Reset</button>
    </div>
    <div class="viewport">
        <div id="img-container">
            <img src="{file_url}" id="target-img">
        </div>
    </div>
    <script>
        let scale = 1;
        let rotation = 0;
        const container = document.getElementById('img-container');
        
        function update() {{
            container.style.transform = `scale(${{scale}}) rotate(${{rotation}}deg)`;
        }}
        
        function zoomIn() {{
            scale *= 1.2;
            update();
        }}
        
        function zoomOut() {{
            scale /= 1.2;
            update();
        }}
        
        function rotateLeft() {{
            rotation -= 90;
            update();
        }}
        
        function rotateRight() {{
            rotation += 90;
            update();
        }}
        
        function reset() {{
            scale = 1;
            rotation = 0;
            update();
        }}
        
        // Mouse wheel zoom
        document.querySelector('.viewport').addEventListener('wheel', (e) => {{
            if (e.ctrlKey) {{
                e.preventDefault();
                if (e.deltaY < 0) zoomIn();
                else zoomOut();
            }}
        }});
    </script>
</body>
</html>"""
            self.web_view.setHtml(html, QUrl.fromLocalFile(os.path.abspath(filepath)))
        except Exception as e:
            self._show_error(str(e), filepath)

    def _load_text_document(self, filepath):
        """Load text document with nice styling."""
        try:
            with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()
            
            # Escape HTML characters
            import html as html_lib
            safe_content = html_lib.escape(content)
            
            html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>Text Viewer</title>
<style>
    body {{
        margin: 0;
        padding: 0;
        background: #fdfbf7;
        color: #2d2926;
        font-family: 'Consolas', 'Monaco', 'Courier New', monospace;
        line-height: 1.5;
        height: 100vh;
        display: flex;
        flex-direction: column;
    }}
    body[data-theme='dark'] {{
        background: #1a1816;
        color: #e8e2d6;
    }}
    .container {{
        flex: 1;
        overflow: auto;
        padding: 20px 40px;
    }}
    pre {{
        margin: 0;
        white-space: pre-wrap;
        word-wrap: break-word;
        font-size: 14px;
        tab-size: 4;
    }}
    .info-bar {{
        padding: 8px 20px;
        background: rgba(0,0,0,0.05);
        border-bottom: 1px solid rgba(0,0,0,0.1);
        font-family: system-ui, sans-serif;
        font-size: 12px;
        color: #666;
        display: flex;
        justify-content: space-between;
    }}
    body[data-theme='dark'] .info-bar {{
        background: rgba(255,255,255,0.05);
        border-bottom: 1px solid rgba(255,255,255,0.1);
        color: #aaa;
    }}
</style>
</head>
<body data-theme="{self._theme}">
    <div class="info-bar">
        <span>{html_lib.escape(os.path.basename(filepath))}</span>
        <span>{len(content)} chars</span>
    </div>
    <div class="container">
        <pre>{safe_content}</pre>
    </div>
</body>
</html>"""
            self.web_view.setHtml(html, QUrl.fromLocalFile(os.path.abspath(filepath)))
        except Exception as e:
            self._show_error(str(e), filepath)
    
    def _show_error(self, message, filepath):
        """Show error message."""
        error_html = f'''<!DOCTYPE html>
<html>
<head><meta charset="utf-8"><title>Error</title></head>
<body style="font-family: sans-serif; padding: 40px; text-align: center; background: #1a1a1a; color: #e0e0e0;">
    <h2 style="color: #ef4444;">Failed to load document</h2>
    <p style="color: #888;">{message}</p>
    <p style="font-size: 12px; color: #666;">{filepath}</p>
</body>
</html>'''
        self.web_view.setHtml(error_html)


# --- On-Screen Keyboard Window ---
def _set_titlebar_color(window, light_color=0xF0F3F5, dark_color=0x1A1A1A):
    """
    Set the Windows 11 titlebar caption colour via DWM.
    DWMWA_CAPTION_COLOR = 35  (Windows 11 build ≥ 22000 only).
    Silently does nothing on older Windows or non-Windows platforms.
    Color value is a COLORREF: 0x00BBGGRR  (little-endian byte order).
    """
    try:
        import ctypes
        # Detect whether the window is in dark mode by checking its theme
        try:
            prefs = load_preferences() or {}
            mode = str(prefs.get('theme_mode') or 'system').lower()
            if mode == 'system':
                try:
                    import winreg
                    key = winreg.OpenKey(winreg.HKEY_CURRENT_USER,
                        r'Software\Microsoft\Windows\CurrentVersion\Themes\Personalize')
                    val, _ = winreg.QueryValueEx(key, 'AppsUseLightTheme')
                    is_dark = not bool(val)
                except Exception:
                    is_dark = False
            else:
                is_dark = (mode == 'dark')
        except Exception:
            is_dark = False

        color = dark_color if is_dark else light_color
        hwnd = int(window.winId())
        # DWMWA_CAPTION_COLOR = 35
        ctypes.windll.dwmapi.DwmSetWindowAttribute(
            hwnd, 35,
            ctypes.byref(ctypes.c_int(color)),
            ctypes.sizeof(ctypes.c_int)
        )
    except Exception:
        pass  # Not Win11, or DWM not available — no-op


class PopupDialog(QMainWindow):
    """
    A floating dialog window that hosts a QWebEngineView pointing at one of the
    /dialog/* Flask pages.  Unlike the keyboard, this window ACCEPTS focus so the
    user can interact with form controls.
    """

    def __init__(self, title: str, url: str, width: int, height: int,
                 server_url: str, parent=None, min_width: int = 400,
                 min_height: int = 300, window_flags=None):
        # Default to a tool window (thin title bar, excluded from taskbar).
        # Individual dialogs can override the flags when they need a full window.
        flags = window_flags if window_flags is not None else Qt.WindowType.Tool
        super().__init__(parent, flags)
        self.server_url = server_url
        self.setWindowTitle(title)
        self.resize(width, height)
        self.setMinimumSize(min_width, min_height)

        if os.path.exists(ICON_PATH):
            self.setWindowIcon(QIcon(ICON_PATH))
        elif os.path.exists(LOGO_PATH):
            self.setWindowIcon(QIcon(LOGO_PATH))

        central = QWidget()
        self.setCentralWidget(central)
        layout = QVBoxLayout(central)
        layout.setContentsMargins(0, 0, 0, 0)

        self.web_view = QWebEngineView()
        self.web_view.setContextMenuPolicy(Qt.ContextMenuPolicy.NoContextMenu)
        settings = self.web_view.settings()
        settings.setAttribute(
            QWebEngineSettings.WebAttribute.LocalContentCanAccessRemoteUrls, True
        )

        # Wire popup pages into the same JS bridge used by the main window.
        # Without this, pyqt-bridge.js cannot initialize in popup dialogs.
        self.channel = QWebChannel()
        bridge_obj = getattr(parent, 'bridge', None)
        if bridge_obj is None:
            bridge_obj = JsBridge(parent or self)
        self.channel.registerObject('pyqt', bridge_obj)
        self.web_view.page().setWebChannel(self.channel)

        layout.addWidget(self.web_view)
        self.web_view.setUrl(QUrl(url))

    def showEvent(self, event):
        super().showEvent(event)
        # Apply titlebar colour after the window has an HWND
        _set_titlebar_color(self)

    def closeEvent(self, event):
        # Hide rather than destroy — avoids QWebEngine teardown cost on re-open
        event.ignore()
        self.hide()


class BlockingPopupDialog(QMainWindow):
    """
    A popup dialog for blocking operations (file import, export, agent run, etc.).

    Differences from PopupDialog:
    - No OS close button (CustomizeWindowHint + WindowTitleHint without WindowCloseButtonHint)
    - closeEvent is a hard no-op — the Python code that started the operation is the
      only thing allowed to hide this window.
    """

    def __init__(self, title: str, url: str, width: int, height: int,
                 server_url: str, parent=None):
        super().__init__(parent,
                 Qt.WindowType.Tool |
                 Qt.WindowType.CustomizeWindowHint |
                 Qt.WindowType.WindowTitleHint)
        self.server_url = server_url
        self.setWindowTitle(title)
        self.resize(width, height)
        self.setMinimumSize(280, 160)
        self.setMaximumSize(width, height)  # prevent resize

        if os.path.exists(ICON_PATH):
            self.setWindowIcon(QIcon(ICON_PATH))
        elif os.path.exists(LOGO_PATH):
            self.setWindowIcon(QIcon(LOGO_PATH))

        central = QWidget()
        self.setCentralWidget(central)
        layout = QVBoxLayout(central)
        layout.setContentsMargins(0, 0, 0, 0)

        self.web_view = QWebEngineView()
        self.web_view.setContextMenuPolicy(Qt.ContextMenuPolicy.NoContextMenu)
        settings = self.web_view.settings()
        settings.setAttribute(
            QWebEngineSettings.WebAttribute.LocalContentCanAccessRemoteUrls, True
        )
        layout.addWidget(self.web_view)
        self.web_view.setUrl(QUrl(url))

    def showEvent(self, event):
        super().showEvent(event)
        _set_titlebar_color(self)

    def closeEvent(self, event):
        # Blocking dialogs cannot be closed by the user
        event.ignore()


class KeyboardWindow(QMainWindow):
    """
    Floating on-screen Sanskrit keyboard.

    Window flags:
      - WindowDoesNotAcceptFocus: clicking keys NEVER steals focus from the main
        editor, so the Quill cursor position is preserved throughout.
      - Tool: thin title bar, excluded from taskbar.

    IMPORTANT: these flags must be passed to super().__init__() — setting them
    later via setWindowFlags() causes a hide/show cycle on Windows (Qt 6 quirk).
    """

    def __init__(self, server_url: str, parent=None):
        super().__init__(
            parent,
            Qt.WindowType.WindowDoesNotAcceptFocus |
            Qt.WindowType.Tool
        )
        self.server_url = server_url
        self.setWindowTitle('Sanskrit Keyboard — śikṣāmitra')
        self.resize(940, 430)
        self.setMinimumSize(640, 320)

        if os.path.exists(ICON_PATH):
            self.setWindowIcon(QIcon(ICON_PATH))
        elif os.path.exists(LOGO_PATH):
            self.setWindowIcon(QIcon(LOGO_PATH))

        central = QWidget()
        self.setCentralWidget(central)
        layout = QVBoxLayout(central)
        layout.setContentsMargins(0, 0, 0, 0)

        self.web_view = QWebEngineView()
        self.web_view.setContextMenuPolicy(Qt.ContextMenuPolicy.NoContextMenu)
        settings = self.web_view.settings()
        settings.setAttribute(
            QWebEngineSettings.WebAttribute.LocalContentCanAccessRemoteUrls, True
        )
        layout.addWidget(self.web_view)

        self.web_view.setUrl(QUrl(f'{server_url}/keyboard'))

    def showEvent(self, event):
        super().showEvent(event)
        _set_titlebar_color(self)

    def closeEvent(self, event):
        # Hide rather than destroy — avoids QWebEngine teardown cost on re-open.
        event.ignore()
        self.hide()


# --- Main Window ---
class MainWindow(QMainWindow):
    """Main application window."""
    
    def __init__(self, server_url):
        super().__init__()
        self.server_url = server_url
        self.viewer_windows = []
        self._close_confirmed = False  # Track if close was confirmed by JS
        
        # Window setup (do NOT show yet — show after fully constructed)
        self.setWindowTitle('śikṣāmitra')
        self.resize(1200, 800)
        
        # Set icon
        if os.path.exists(ICON_PATH):
            self.setWindowIcon(QIcon(ICON_PATH))
        elif os.path.exists(LOGO_PATH):
            self.setWindowIcon(QIcon(LOGO_PATH))
        
        # Central widget
        central = QWidget()
        self.setCentralWidget(central)
        layout = QVBoxLayout(central)
        layout.setContentsMargins(0, 0, 0, 0)
        
        # WebEngine View
        self.web_view = QWebEngineView()
        self.web_view.setContextMenuPolicy(Qt.ContextMenuPolicy.NoContextMenu)
        layout.addWidget(self.web_view)
        
        # Setup web channel for JS bridge
        self.channel = QWebChannel()
        self.bridge = JsBridge(self)
        self.channel.registerObject('pyqt', self.bridge)
        self.web_view.page().setWebChannel(self.channel)
        
        # Set page background to match theme to prevent flash during load
        try:
            prefs = load_preferences()
            theme_mode = str(prefs.get('theme_mode', 'light')).lower()
            if theme_mode == 'dark':
                bg_color = QColor(17, 17, 19)   # #111113
            elif theme_mode == 'system':
                try:
                    import winreg
                    key = winreg.OpenKey(winreg.HKEY_CURRENT_USER,
                        r'Software\Microsoft\Windows\CurrentVersion\Themes\Personalize')
                    val, _ = winreg.QueryValueEx(key, 'AppsUseLightTheme')
                    winreg.CloseKey(key)
                    bg_color = QColor(17, 17, 19) if val == 0 else QColor(242, 242, 245)
                except Exception:
                    bg_color = QColor(242, 242, 245)
            else:
                bg_color = QColor(242, 242, 245)  # #f2f2f5
            self.web_view.page().setBackgroundColor(bg_color)
        except Exception:
            pass
        
        # Connect signals
        self.bridge.viewerWindowRequested.connect(self._open_viewer_window)
        self.bridge.closeConfirmed.connect(self._on_close_confirmed)
        self.bridge.keyboardToggleRequested.connect(self._toggle_keyboard)
        self.bridge.openDialogRequested.connect(self._open_dialog)
        self.bridge.closeDialogRequested.connect(self._close_dialog)
        self.bridge.showLoaderRequested.connect(self._show_loader_dialog)
        self.bridge.hideLoaderRequested.connect(self._hide_loader_dialog)

        # On-screen keyboard window (created lazily on first toggle)
        self._keyboard_window = None
        # Popup dialog windows, keyed by dialog_id
        self._dialog_windows = {}
        # Blocking loader OS window (created lazily on first use)
        self._loader_window = None

        # Load app
        self.web_view.setUrl(QUrl(server_url))
    
    def _open_viewer_window(self, filepath, title):
        """Open a new viewer window."""
        if not title:
            title = os.path.splitext(os.path.basename(filepath))[0]
        
        # Get current theme from preferences
        theme = 'light'
        try:
            prefs = load_preferences()
            theme_mode = str(prefs.get('theme_mode', 'light')).lower()
            if theme_mode == 'dark':
                theme = 'dark'
            elif theme_mode == 'system':
                # On Windows, check if dark mode is enabled
                try:
                    import winreg
                    key = winreg.OpenKey(winreg.HKEY_CURRENT_USER, 
                        r"Software\Microsoft\Windows\CurrentVersion\Themes\Personalize")
                    value, _ = winreg.QueryValueEx(key, "AppsUseLightTheme")
                    theme = 'light' if value else 'dark'
                except Exception:
                    theme = 'light'
        except Exception:
            theme = 'light'
        
        viewer = ViewerWindow(filepath, title, self.server_url, theme, None)
        viewer.show()
        self.viewer_windows.append(viewer)
    
    def _on_close_confirmed(self, can_close):
        """Called from JS after unsaved changes check."""
        if can_close:
            self._close_confirmed = True
            self.close()  # Re-trigger close, this time it will succeed

    def _toggle_keyboard(self):
        """Toggle the floating on-screen Sanskrit keyboard."""
        if self._keyboard_window is None:
            self._keyboard_window = KeyboardWindow(self.server_url, parent=None)
            # Position: centred horizontally above the main window's bottom edge
            geo = self.geometry()
            kw, kh = 980, 340
            kx = geo.left() + max(0, (geo.width() - kw) // 2)
            ky = max(0, geo.bottom() - kh - 40)
            self._keyboard_window.setGeometry(kx, ky, kw, kh)

        if self._keyboard_window.isVisible():
            self._keyboard_window.hide()
        else:
            self._keyboard_window.show()
            # Immediately reclaim focus for the main editor.
            # WindowDoesNotAcceptFocus prevents most focus theft, but an
            # explicit setFocus() after show() ensures the Quill cursor stays visible.
            self.web_view.setFocus()

    def _open_dialog(self, dialog_id: str):
        """Open (or raise) a popup dialog window by its ID."""
        DIALOG_SPECS = {
            'autorun':         ('Run Agent — śikṣāmitra',             '/dialog/autorun',         720, 520),
            'shloka':          ('Insert Śloka — śikṣāmitra',          '/dialog/shloka',          800, 560),
            'iast':            ('IAST Characters — śikṣāmitra',       '/dialog/iast',            700, 480),
            'autosvara':       ('Automatic Svaras — śikṣāmitra',      '/dialog/autosvara',       620, 520),
            'paragraphstyles': ('Paragraph Styles — śikṣāmitra',      '/dialog/paragraphstyles', 760, 620),
            'media':           ('Embedded Media — śikṣāmitra',        '/dialog/media',           680, 480),
            'audio-picker':    ('Attach Audio — śikṣāmitra',           '/dialog/audio-picker',    560, 520),
            'audio-editor':    ('Audio Editor — śikṣāmitra',           '/dialog/audio-editor',    1100, 680),
            'message':         ('śikṣāmitra',                          '/dialog/message',         500, 220),
        }
        if dialog_id not in DIALOG_SPECS:
            return

        title, path, w, h = DIALOG_SPECS[dialog_id]
        url = f'{self.server_url}{path}'
        min_w, min_h = (360, 190) if dialog_id == 'message' else (400, 300)

        # Stateful dialogs (audio picker / editor / loader / media / message) must
        # reload their URL on each open so they re-fetch the current state from the server.
        STATEFUL = {'audio-picker', 'audio-editor', 'media', 'message'}

        win = self._dialog_windows.get(dialog_id)
        if win is None:
            dialog_flags = None
            dialog_parent = self
            if dialog_id == 'audio-editor':
                dialog_flags = (
                    Qt.WindowType.Window |
                    Qt.WindowType.WindowTitleHint |
                    Qt.WindowType.WindowMinimizeButtonHint |
                    Qt.WindowType.WindowMaximizeButtonHint |
                    Qt.WindowType.WindowCloseButtonHint
                )
            win = PopupDialog(title, url, w, h, self.server_url, parent=dialog_parent,
                              min_width=min_w, min_height=min_h, window_flags=dialog_flags)
            self._dialog_windows[dialog_id] = win
            # Centre over the main window
            geo = self.geometry()
            dx = geo.left() + (geo.width()  - w) // 2
            dy = geo.top()  + (geo.height() - h) // 2
            win.setGeometry(dx, dy, w, h)
        elif dialog_id in STATEFUL:
            # Force the page to reload so it picks up freshly staged state
            win.web_view.setUrl(QUrl(url))

        if win.isVisible():
            win.raise_()
            win.activateWindow()
        else:
            win.show()

    def _close_dialog(self, dialog_id: str):
        """Hide a popup dialog window (called from JS via bridge.closeDialog)."""
        win = self._dialog_windows.get(dialog_id)
        if win is not None and win.isVisible():
            win.hide()

    def _show_loader_dialog(self, title: str, message: str):
        """Show the blocking loader OS window for long-running operations."""
        global _loader_state
        with _loader_state_lock:
            _loader_state = {'title': title or 'Loading\u2026', 'message': message or ''}

        url = f'{self.server_url}/dialog/loader'
        if self._loader_window is None:
            self._loader_window = BlockingPopupDialog(
                'śikṣāmitra', url, 340, 200, self.server_url, parent=self
            )
        else:
            # Navigate to force a fresh poll of state after update
            self._loader_window.web_view.setUrl(QUrl(url))

        # Centre over the main window
        geo = self.geometry()
        lw, lh = 340, 200
        lx = geo.left() + (geo.width()  - lw) // 2
        ly = geo.top()  + (geo.height() - lh) // 2
        self._loader_window.setGeometry(lx, ly, lw, lh)
        self._loader_window.show()
        self._loader_window.raise_()

    def _hide_loader_dialog(self):
        """Hide the blocking loader OS window."""
        if self._loader_window is not None:
            self._loader_window.hide()

    def closeEvent(self, event):
        """Handle close event with unsaved changes confirmation."""
        # If already confirmed by JS, allow close
        if self._close_confirmed:
            for v in self.viewer_windows:
                try:
                    v.close()
                except Exception:
                    pass
            # Close keyboard window if open
            if self._keyboard_window is not None:
                try:
                    self._keyboard_window.web_view.setUrl(QUrl('about:blank'))
                    self._keyboard_window.destroy()
                except Exception:
                    pass
                self._keyboard_window = None
            # Close all popup dialog windows
            for dlg in list(self._dialog_windows.values()):
                try:
                    dlg.web_view.setUrl(QUrl('about:blank'))
                    dlg.destroy()
                except Exception:
                    pass
            self._dialog_windows.clear()
            # Close loader window if open
            if self._loader_window is not None:
                try:
                    self._loader_window.web_view.setUrl(QUrl('about:blank'))
                    self._loader_window.destroy()
                except Exception:
                    pass
                self._loader_window = None
            event.accept()
            return
        
        # Ask JS if we can close (checks unsaved changes)
        event.ignore()
        self.web_view.page().runJavaScript(
            """
            (async function() {
                try {
                    if (typeof window.__sikshamitra_canClose === 'function') {
                        const canClose = await window.__sikshamitra_canClose();
                        if (window.pywebview && window.pywebview.api && typeof window.pywebview.api._confirm_close === 'function') {
                            window.pywebview.api._confirm_close(!!canClose);
                        }
                        return canClose;
                    }
                    // No function defined, allow close
                    if (window.pywebview && window.pywebview.api && typeof window.pywebview.api._confirm_close === 'function') {
                        window.pywebview.api._confirm_close(true);
                    }
                    return true;
                } catch (e) {
                    console.error('Close check error:', e);
                    if (window.pywebview && window.pywebview.api && typeof window.pywebview.api._confirm_close === 'function') {
                        window.pywebview.api._confirm_close(true);
                    }
                    return true;
                }
            })();
            """
        )


# --- Splash Screen ---
class SplashScreen(QSplashScreen):
    """Loading splash screen."""
    
    def __init__(self):
        # Create pixmap with logo
        if os.path.exists(LOGO_PATH):
            pixmap = QPixmap(LOGO_PATH).scaled(
                120, 120, Qt.AspectRatioMode.KeepAspectRatio,
                Qt.TransformationMode.SmoothTransformation
            )
        else:
            pixmap = QPixmap(200, 200)
            pixmap.fill(Qt.GlobalColor.transparent)
        
        super().__init__(pixmap)
        self.setWindowFlags(
            Qt.WindowType.SplashScreen |
            Qt.WindowType.FramelessWindowHint |
            Qt.WindowType.WindowStaysOnTopHint
        )
    
    def showMessage(self, message):
        super().showMessage(
            message,
            Qt.AlignmentFlag.AlignBottom | Qt.AlignmentFlag.AlignHCenter,
            Qt.GlobalColor.gray
        )


# --- Main Entry Point ---
def main():
    # Create Qt application
    qt_app = QApplication(sys.argv)
    qt_app.setApplicationName('śikṣāmitra')
    qt_app.setOrganizationName('śikṣāmitra')
    
    # Set application icon
    if os.path.exists(ICON_PATH):
        qt_app.setWindowIcon(QIcon(ICON_PATH))
    elif os.path.exists(LOGO_PATH):
        qt_app.setWindowIcon(QIcon(LOGO_PATH))
    
    # Start Flask server in background
    server = FlaskServerThread(app, host='127.0.0.1', port=0)
    server.start()
    
    if not wait_for_server(server.host, server.port, timeout=10.0):
        QMessageBox.critical(None, "Error", "Failed to start server")
        return 1
    
    server_url = f'http://{server.host}:{server.port}'
    
    # Create and show main window (single show after full construction)
    main_window = MainWindow(server_url)
    main_window.showMaximized()
    
    # Run event loop
    result = qt_app.exec()
    
    # Cleanup
    server.shutdown()
    return result


# ─── On-screen keyboard routes ───────────────────────────────────────────────
# These must appear before the catch-all serve_static route.

@app.route('/keyboard')
def serve_keyboard():
    """Serve the on-screen Sanskrit keyboard page."""
    keyboard_path = os.path.join(BASE_DIR, 'keyboard.html')
    try:
        with open(keyboard_path, 'r', encoding='utf-8') as f:
            html = f.read()
        return Response(html, mimetype='text/html')
    except FileNotFoundError:
        return Response('<h1>keyboard.html not found</h1>', status=404, mimetype='text/html')


@app.route('/api/keyboard/insert', methods=['POST'])
def keyboard_insert():
    """Receive a character from the on-screen keyboard and queue it for the editor."""
    data = request.get_json(force=True, silent=True) or {}
    char = data.get('char', '')
    if char:
        with _keyboard_queue_lock:
            _keyboard_queue.append(char)
    return jsonify({'status': 'ok'})


@app.route('/api/keyboard/poll', methods=['GET'])
def keyboard_poll():
    """Return and drain all pending keyboard characters for the editor to insert."""
    with _keyboard_queue_lock:
        chars = list(_keyboard_queue)
        _keyboard_queue.clear()
    return jsonify({'chars': chars})


# ─── Popup dialog pages ──────────────────────────────────────────────────────

def _serve_dialog_page(filename):
    """Helper: read a dialog HTML, inject theme and base-href, return Response."""
    path = os.path.join(BASE_DIR, filename)
    try:
        with open(path, 'r', encoding='utf-8') as f:
            html = f.read()
        # Inject <base href="/"> so relative asset paths (styles.css, logo.png, …)
        # resolve correctly even though the page is served under /dialog/…
        html = html.replace('<head>', '<head><base href="/">', 1)
        # Inject theme the same way the main editor does
        try:
            prefs = load_preferences() or {}
            theme_mode = str(prefs.get('theme_mode') or 'system').lower()
            if theme_mode not in ('light', 'dark', 'system'):
                theme_mode = 'system'
        except Exception:
            theme_mode = 'system'
        html = html.replace('__SIKSAMITRA_INITIAL_THEME_MODE__', theme_mode)
        return Response(html, mimetype='text/html')
    except FileNotFoundError:
        return Response(f'<h1>{filename} not found</h1>', status=404, mimetype='text/html')


@app.route('/dialog/autorun')
def serve_dialog_autorun():
    return _serve_dialog_page('dialog-autorun.html')


@app.route('/dialog/shloka')
def serve_dialog_shloka():
    return _serve_dialog_page('dialog-shloka.html')


@app.route('/dialog/iast')
def serve_dialog_iast():
    return _serve_dialog_page('dialog-iast.html')


@app.route('/dialog/autosvara')
def serve_dialog_autosvara():
    return _serve_dialog_page('dialog-autosvara.html')


@app.route('/dialog/paragraphstyles')
def serve_dialog_paragraphstyles():
    return _serve_dialog_page('dialog-paragraphstyles.html')


@app.route('/dialog/media')
def serve_dialog_media():
    return _serve_dialog_page('dialog-media.html')


@app.route('/dialog/audio-picker')
def serve_dialog_audio_picker():
    return _serve_dialog_page('dialog-audio-picker.html')


@app.route('/dialog/audio-editor')
def serve_dialog_audio_editor():
    return _serve_dialog_page('dialog-audio-editor.html')


@app.route('/dialog/message')
def serve_dialog_message():
    return _serve_dialog_page('dialog-message.html')


# ── Modal-message state — editor stages config, popup reads, posts response ──
_message_state = {}
_message_state_lock = threading.Lock()


@app.route('/api/message/state', methods=['GET', 'POST'])
def message_state():
    """GET: dialog-message reads its current config.
    POST: editor sets the next message to display (id, type, title, message, buttons, ...)."""
    global _message_state
    if request.method == 'POST':
        data = request.get_json(force=True, silent=True) or {}
        with _message_state_lock:
            _message_state = dict(data)
        return jsonify({'status': 'ok'})
    with _message_state_lock:
        return jsonify(dict(_message_state))


# ── Audio dialog state — editor writes before opening popup, popup reads/writes ──
# Two separate buckets (picker and editor) so one dialog can't clobber the other.
_audio_picker_state = {}       # Data for dialog-audio-picker (audio library, target lines)
_audio_editor_state = {}       # Data for dialog-audio-editor (audio data, regions, targets)
_audio_sections_cache = {'sections': []}  # All available sections for the "+ Section" picker
_audio_state_lock   = threading.Lock()


@app.route('/api/audio/editor/sections', methods=['GET', 'POST'])
def audio_editor_sections():
    """GET: audio-editor's section-picker reads available sections.
    POST: editor writes the cached list (called whenever the editor opens an audio dialog)."""
    global _audio_sections_cache
    if request.method == 'POST':
        data = request.get_json(force=True, silent=True) or {}
        with _audio_state_lock:
            _audio_sections_cache = {'sections': data.get('sections', [])}
        return jsonify({'status': 'ok'})
    with _audio_state_lock:
        return jsonify(dict(_audio_sections_cache))


@app.route('/api/audio/picker/state', methods=['GET', 'POST'])
def audio_picker_state():
    """GET: audio-picker popup reads the currently-staged data.
    POST: editor writes the data before opening the popup."""
    global _audio_picker_state
    if request.method == 'POST':
        data = request.get_json(force=True, silent=True) or {}
        with _audio_state_lock:
            _audio_picker_state = dict(data)
        return jsonify({'status': 'ok'})
    with _audio_state_lock:
        return jsonify(dict(_audio_picker_state))


@app.route('/api/audio/editor/state', methods=['GET', 'POST'])
def audio_editor_state():
    """GET: audio-editor popup reads the currently-staged data.
    POST: editor writes the data before opening the popup."""
    global _audio_editor_state
    if request.method == 'POST':
        data = request.get_json(force=True, silent=True) or {}
        with _audio_state_lock:
            _audio_editor_state = dict(data)
        return jsonify({'status': 'ok'})
    with _audio_state_lock:
        return jsonify(dict(_audio_editor_state))


# ── Embedded audio cache — written by main editor before opening media popup ──
_embedded_audio_cache = []
_embedded_audio_lock = threading.Lock()


@app.route('/api/audio/embedded', methods=['GET', 'POST'])
def audio_embedded():
    """Cache the list of embedded audio items (metadata only) for the media popup."""
    global _embedded_audio_cache
    if request.method == 'POST':
        data = request.get_json(force=True, silent=True) or {}
        with _embedded_audio_lock:
            _embedded_audio_cache = data.get('audio', [])
        return jsonify({'status': 'ok'})
    with _embedded_audio_lock:
        return jsonify({'audio': list(_embedded_audio_cache)})


@app.route('/api/dialog/action', methods=['POST'])
def dialog_action():
    """Receive an action from a popup dialog and queue it for the main editor."""
    data = request.get_json(force=True, silent=True) or {}
    if data:
        with _dialog_action_lock:
            _dialog_action_queue.append(data)
    return jsonify({'status': 'queued'})


@app.route('/api/dialog/actions', methods=['GET'])
def dialog_actions():
    """Return and drain all pending dialog actions for the main editor."""
    with _dialog_action_lock:
        actions = list(_dialog_action_queue)
        _dialog_action_queue.clear()
    return jsonify({'actions': actions})


@app.route('/dialog/loader')
def serve_dialog_loader():
    return _serve_dialog_page('dialog-loader.html')


# ── Blocking-loader state — updated by the main editor, read by dialog-loader.html ──
_loader_state = {'title': 'Loading\u2026', 'message': ''}
_loader_state_lock = threading.Lock()


def _update_loader(title, message=''):
    """Convenience: update the loader state from any thread (e.g. Flask request handlers).

    The OS loader window polls /api/loader/state every 300ms, so updates appear
    within ~300ms of this call.
    """
    global _loader_state
    with _loader_state_lock:
        _loader_state = {
            'title': str(title or 'Loading\u2026'),
            'message': str(message or ''),
        }


@app.route('/api/loader/state', methods=['GET'])
def loader_state_get():
    """Return current loader title/message for the blocking-loader popup to poll."""
    with _loader_state_lock:
        return jsonify(dict(_loader_state))


@app.route('/api/loader/update', methods=['POST'])
def loader_state_update():
    """Update loader title/message mid-operation (e.g. during multi-step agent run)."""
    global _loader_state
    data = request.get_json(force=True, silent=True) or {}
    with _loader_state_lock:
        if 'title'   in data: _loader_state['title']   = str(data['title']   or '')
        if 'message' in data: _loader_state['message'] = str(data['message'] or '')
    return jsonify({'status': 'ok'})


@app.route('/<path:filename>')
def serve_static(filename):
    """Serve static files."""
    # Explicitly block API routes from being served as static files
    if filename.startswith('api/'):
        return jsonify({'error': 'Not found'}), 404
    return send_from_directory(BASE_DIR, filename)


if __name__ == '__main__':
    sys.exit(main())
